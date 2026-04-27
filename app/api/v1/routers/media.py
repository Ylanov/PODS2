# app/api/v1/routers/media.py
"""
Учёт машинных носителей информации (МНИ): флешки, SSD, HDD, SD-карты, диски.

Эндпоинты:
  GET    /media                          — список носителей с фильтрами + summary
  POST   /media                          — создать
  PUT    /media/{id}                     — изменить
  DELETE /media/{id}                     — удалить
  POST   /media/{id}/issue               — выдать (зафиксировать держателя + журнал)
  POST   /media/{id}/return              — вернуть на хранение (журнал)
  GET    /media/{id}/transfers           — история движений
  GET    /media/tags-export              — выгрузить .docx с сеткой бирок

Доступ:
  • role='unit' — только свой отдел
  • admin       — любой отдел через ?unit=<username>
  • остальные   — 403
"""

import io
from datetime import datetime, date as date_type
from typing import Literal, Optional
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, ConfigDict, Field
from sqlalchemy import or_
from sqlalchemy.orm import Session

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.db.database import get_db
from app.models.user import User
from app.models.media import (
    MediaItem, MediaTransfer,
    MEDIA_TYPES, MEDIA_CLASSIFICATIONS, MEDIA_STATUSES, TRANSFER_KINDS,
)
from app.api.dependencies import get_current_user


router = APIRouter()


# ─── Схемы ──────────────────────────────────────────────────────────────────

class MediaIn(BaseModel):
    inv_number:        str = Field(..., min_length=1, max_length=50)
    media_type:        Literal["flash", "ssd", "hdd", "sd", "cd_dvd", "other"] = "flash"
    serial_number:     Optional[str] = Field(None, max_length=120)
    capacity_gb:       Optional[int] = Field(None, ge=0, le=100_000)
    classification:    Literal["open", "dsp", "secret", "top_secret"] = "dsp"
    status:            Literal["available", "issued", "broken",
                               "written_off", "lost"] = "available"
    # Связь с общей базой людей. Если задана, при чтении ФИО/подразделение
    # подтягиваются из Person, а не из кеш-полей ниже.
    holder_person_id:  Optional[int] = None
    holder_full_name:  Optional[str] = Field(None, max_length=300)
    holder_short_name: Optional[str] = Field(None, max_length=100)
    holder_department: Optional[str] = Field(None, max_length=100)
    issue_date:        Optional[date_type] = None
    last_check_date:   Optional[date_type] = None
    next_check_date:   Optional[date_type] = None
    notes:             Optional[str] = None


class TransferOut(BaseModel):
    id:               int
    kind:             str
    event_date:       date_type
    person_full_name: Optional[str] = None
    department:       Optional[str] = None
    operator:         Optional[str] = None
    notes:            Optional[str] = None
    created_at:       datetime

    model_config = ConfigDict(from_attributes=True)


class MediaOut(BaseModel):
    id:                int
    unit_username:     str
    inv_number:        str
    media_type:        str
    serial_number:     Optional[str] = None
    capacity_gb:       Optional[int] = None
    classification:    str
    status:            str
    holder_person_id:  Optional[int] = None
    holder_full_name:  Optional[str] = None
    holder_short_name: Optional[str] = None
    holder_department: Optional[str] = None
    issue_date:        Optional[date_type] = None
    last_check_date:   Optional[date_type] = None
    next_check_date:   Optional[date_type] = None
    notes:             Optional[str]      = None
    created_at:        datetime
    updated_at:        datetime

    model_config = ConfigDict(from_attributes=True)


def _media_to_out(item: MediaItem) -> MediaOut:
    """
    Преобразование MediaItem → MediaOut с подменой ФИО/подразделения
    значениями из связанной Person (если есть). Это даёт «онлайн»-эффект:
    переименовали человека в общей базе → в учёте МНИ сразу новое имя.
    """
    p = item.holder_person
    full = (p.full_name  if p else item.holder_full_name)
    dept = (p.department if p else item.holder_department)
    return MediaOut(
        id=item.id,
        unit_username=item.unit_username,
        inv_number=item.inv_number,
        media_type=item.media_type,
        serial_number=item.serial_number,
        capacity_gb=item.capacity_gb,
        classification=item.classification,
        status=item.status,
        holder_person_id=item.holder_person_id,
        holder_full_name=full,
        holder_short_name=item.holder_short_name,
        holder_department=dept,
        issue_date=item.issue_date,
        last_check_date=item.last_check_date,
        next_check_date=item.next_check_date,
        notes=item.notes,
        created_at=item.created_at,
        updated_at=item.updated_at,
    )


class IssueRequest(BaseModel):
    """Тело /issue — выдача носителя."""
    holder_full_name:  str = Field(..., min_length=1, max_length=300)
    holder_short_name: Optional[str] = Field(None, max_length=100)
    holder_department: Optional[str] = Field(None, max_length=100)
    issue_date:        date_type
    notes:             Optional[str] = None


class MediaSummary(BaseModel):
    total:        int
    available:    int
    issued:       int
    broken:       int
    written_off:  int
    lost:         int
    by_type:      dict[str, int]
    by_class:     dict[str, int]
    by_department: dict[str, int]
    overdue_check: int   # просрочена очередная проверка


class MediaListResponse(BaseModel):
    unit_username: str
    items:         list[MediaOut]
    summary:       MediaSummary


# ─── Вспомогательные ────────────────────────────────────────────────────────

def _resolve_unit(current_user: User, unit_override: Optional[str]) -> str:
    if current_user.role == "admin":
        return (unit_override or current_user.username).strip()
    if current_user.role == "unit":
        return current_user.username
    raise HTTPException(status_code=403, detail="Доступ только для отдела или админа")


def _check_unit(item: MediaItem, current_user: User) -> None:
    if not item:
        raise HTTPException(status_code=404, detail="Носитель не найден")
    if current_user.role == "unit" and item.unit_username != current_user.username:
        raise HTTPException(status_code=403, detail="Чужой отдел")
    if current_user.role not in ("admin", "unit"):
        raise HTTPException(status_code=403, detail="Доступ запрещён")


def _compute_summary(items: list[MediaItem]) -> MediaSummary:
    today = date_type.today()
    by_type:  dict[str, int] = {t: 0 for t in MEDIA_TYPES}
    by_class: dict[str, int] = {c: 0 for c in MEDIA_CLASSIFICATIONS}
    by_dept:  dict[str, int] = {}
    counts = {"available": 0, "issued": 0, "broken": 0,
              "written_off": 0, "lost": 0}
    overdue = 0
    for it in items:
        counts[it.status] = counts.get(it.status, 0) + 1
        by_type[it.media_type]      = by_type.get(it.media_type, 0) + 1
        by_class[it.classification] = by_class.get(it.classification, 0) + 1
        if it.holder_department:
            by_dept[it.holder_department] = by_dept.get(it.holder_department, 0) + 1
        if it.next_check_date and it.next_check_date < today:
            overdue += 1
    return MediaSummary(
        total         = len(items),
        available     = counts["available"],
        issued        = counts["issued"],
        broken        = counts["broken"],
        written_off   = counts["written_off"],
        lost          = counts["lost"],
        by_type       = by_type,
        by_class      = by_class,
        by_department = by_dept,
        overdue_check = overdue,
    )


# ─── GET: список + summary ──────────────────────────────────────────────────

@router.get("", response_model=MediaListResponse)
def list_media(
        unit:           Optional[str] = Query(None, max_length=100),
        q:              Optional[str] = Query(None, description="Поиск по инв.№ / ФИО / серийному"),
        status_filter:  Optional[str] = Query(None, alias="status"),
        type_filter:    Optional[str] = Query(None, alias="type"),
        dept_filter:    Optional[str] = Query(None, alias="dept",
                                              description="Фильтр по подразделению (точное совпадение)"),
        db:             Session       = Depends(get_db),
        current_user:   User          = Depends(get_current_user),
):
    unit_name = _resolve_unit(current_user, unit)
    query = db.query(MediaItem).filter(MediaItem.unit_username == unit_name)

    if status_filter:
        query = query.filter(MediaItem.status == status_filter)
    if type_filter:
        query = query.filter(MediaItem.media_type == type_filter)
    if dept_filter:
        # Спец-значение «—» (тире) — без подразделения (NULL или пусто)
        if dept_filter in ("—", "-", "_none_"):
            query = query.filter(or_(
                MediaItem.holder_department.is_(None),
                MediaItem.holder_department == "",
            ))
        else:
            query = query.filter(MediaItem.holder_department == dept_filter)
    if q:
        pattern = f"%{q.strip()}%"
        query = query.filter(or_(
            MediaItem.inv_number.ilike(pattern),
            MediaItem.holder_full_name.ilike(pattern),
            MediaItem.holder_short_name.ilike(pattern),
            MediaItem.serial_number.ilike(pattern),
        ))

    items = query.order_by(MediaItem.inv_number.asc()).all()

    return MediaListResponse(
        unit_username = unit_name,
        items         = [_media_to_out(i) for i in items],
        summary       = _compute_summary(items),
    )


# ─── CRUD ───────────────────────────────────────────────────────────────────

def _apply_payload(item: MediaItem, payload: MediaIn) -> None:
    for f in (
        "inv_number", "media_type", "serial_number",
        "capacity_gb", "classification", "status",
        "holder_person_id",
        "holder_full_name", "holder_short_name",
        "holder_department", "issue_date",
        "last_check_date", "next_check_date", "notes",
    ):
        setattr(item, f, getattr(payload, f))


@router.post("", response_model=MediaOut, status_code=201)
def create_media(
        payload:      MediaIn,
        unit:         Optional[str] = Query(None, max_length=100),
        db:           Session       = Depends(get_db),
        current_user: User          = Depends(get_current_user),
):
    unit_name = _resolve_unit(current_user, unit)
    # Проверка уникальности инв. номера в рамках отдела
    exists = (
        db.query(MediaItem.id)
          .filter(MediaItem.unit_username == unit_name,
                  MediaItem.inv_number    == payload.inv_number)
          .first()
    )
    if exists:
        raise HTTPException(
            status_code=409,
            detail=f"Носитель с инв. № «{payload.inv_number}» уже существует",
        )

    item = MediaItem(unit_username=unit_name)
    _apply_payload(item, payload)
    db.add(item)
    db.commit()
    db.refresh(item)
    return _media_to_out(item)


@router.put("/{item_id}", response_model=MediaOut)
def update_media(
        item_id:      int,
        payload:      MediaIn,
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    item = db.query(MediaItem).filter(MediaItem.id == item_id).first()
    _check_unit(item, current_user)
    _apply_payload(item, payload)
    db.commit()
    db.refresh(item)
    return _media_to_out(item)


# ─── Bulk-очистка ───────────────────────────────────────────────────────────
# При массовом импорте шаблона из Excel легко случайно создать сотни пустых
# *-ДСП записей (если в шаблоне 500 строк с предзаполненной нумерацией, а
# реальных данных меньше). Чтобы не удалять руками, есть два режима:
#   • mode=empty (по умолчанию) — сносим записи без держателя/серийника/
#     объёма/примечаний: чисто «пустышки», созданные импортом.
#   • mode=all — полная очистка таблицы отдела (только админ).
# Маршрут декларирован ДО /{item_id}, чтобы буквальный путь /cleanup не
# попал в параметризованный роут.

@router.delete("/cleanup")
def cleanup_media(
        mode:         Literal["empty", "all"] = Query("empty"),
        unit:         Optional[str] = Query(None, max_length=100),
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    unit_name = _resolve_unit(current_user, unit)
    q = db.query(MediaItem).filter(MediaItem.unit_username == unit_name)

    if mode == "empty":
        # «Пустая» = нет ни одного значимого поля кроме инв.№
        q = q.filter(
            or_(MediaItem.holder_full_name.is_(None),  MediaItem.holder_full_name  == ""),
            or_(MediaItem.holder_short_name.is_(None), MediaItem.holder_short_name == ""),
            or_(MediaItem.serial_number.is_(None),     MediaItem.serial_number     == ""),
            MediaItem.capacity_gb.is_(None),
            or_(MediaItem.notes.is_(None),             MediaItem.notes             == ""),
            MediaItem.holder_person_id.is_(None),
        )
    elif mode == "all":
        if current_user.role != "admin":
            raise HTTPException(status_code=403,
                                detail="Полная очистка доступна только админу")

    deleted = q.delete(synchronize_session=False)
    db.commit()
    return {"deleted": deleted, "mode": mode, "unit": unit_name}


@router.delete("/{item_id}", status_code=204)
def delete_media(
        item_id:      int,
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    item = db.query(MediaItem).filter(MediaItem.id == item_id).first()
    _check_unit(item, current_user)
    db.delete(item)
    db.commit()
    return None


class ReassignRequest(BaseModel):
    """Тело /reassign — переписать носитель на другого человека."""
    holder_person_id:  Optional[int] = None
    holder_full_name:  Optional[str] = Field(None, max_length=300)
    holder_short_name: Optional[str] = Field(None, max_length=100)
    holder_department: Optional[str] = Field(None, max_length=100)
    issue_date:        Optional[date_type] = None
    notes:             Optional[str] = None


@router.post("/{item_id}/clear", response_model=MediaOut)
def clear_holder(
        item_id:      int,
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    """
    «Очистка»: убирает держателя — носитель остаётся в учёте, но без
    закрепления (status='available'). Запись о бывшем держателе попадает
    в журнал движений для аудита.
    """
    item = db.query(MediaItem).filter(MediaItem.id == item_id).first()
    _check_unit(item, current_user)

    # Сначала пишем в журнал — забираем текущие значения до очистки
    if item.holder_full_name or item.holder_short_name:
        db.add(MediaTransfer(
            media_id    = item.id,
            kind        = "returned",
            event_date  = date_type.today(),
            person_full_name = item.holder_full_name,
            department  = item.holder_department,
            operator    = current_user.username,
            notes       = "Очистка держателя",
        ))

    item.holder_person_id  = None
    item.holder_full_name  = None
    item.holder_short_name = None
    item.holder_department = None
    item.issue_date        = None
    item.status            = "available"

    db.commit()
    db.refresh(item)
    return _media_to_out(item)


@router.post("/{item_id}/reassign", response_model=MediaOut)
def reassign_holder(
        item_id:      int,
        payload:      ReassignRequest,
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    """
    «Переписать»: меняет держателя носителя. Если задан holder_person_id,
    ФИО и подразделение можно не передавать — подтянутся из общей базы;
    если оператор хочет вписать вручную (редкий случай), может прислать
    holder_full_name/short_name напрямую без person_id.
    """
    item = db.query(MediaItem).filter(MediaItem.id == item_id).first()
    _check_unit(item, current_user)

    # Журнал — кто был раньше
    if item.holder_full_name or item.holder_short_name:
        db.add(MediaTransfer(
            media_id    = item.id,
            kind        = "transferred",
            event_date  = payload.issue_date or date_type.today(),
            person_full_name = item.holder_full_name,
            department  = item.holder_department,
            operator    = current_user.username,
            notes       = "Передача другому держателю" + (
                f": {payload.notes}" if payload.notes else ""),
        ))

    # Если выбрана персона — подтягиваем full_name/department из неё
    full = payload.holder_full_name
    dept = payload.holder_department
    short = payload.holder_short_name
    if payload.holder_person_id:
        from app.models.person import Person
        p = db.query(Person).filter(Person.id == payload.holder_person_id).first()
        if p:
            full = p.full_name
            dept = p.department
            # Краткое — если оператор не задал, генерим из ФИО
            if not short:
                parts = (p.full_name or "").split()
                if len(parts) >= 2:
                    short = parts[0] + " " + "".join(w[0] + "." for w in parts[1:])

    item.holder_person_id  = payload.holder_person_id
    item.holder_full_name  = full
    item.holder_short_name = short
    item.holder_department = dept
    item.issue_date        = payload.issue_date or date_type.today()
    item.last_check_date   = item.issue_date
    item.status            = "issued" if (full or short) else "available"

    db.commit()
    db.refresh(item)
    return _media_to_out(item)


@router.post("/{item_id}/decommission", response_model=MediaOut)
def decommission_media(
        item_id:      int,
        reason:       Optional[str] = Query(None, max_length=500,
                                            description="Причина списания"),
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    """
    Списание носителя. Не удаляет запись (история сохраняется), а ставит
    status='written_off'. Держатель остаётся в полях для аудита: видно,
    у кого была флешка на момент списания.
    """
    item = db.query(MediaItem).filter(MediaItem.id == item_id).first()
    _check_unit(item, current_user)

    item.status = "written_off"
    if reason:
        item.notes = (item.notes + "\n" if item.notes else "") + f"[Списан] {reason}"
    db.add(MediaTransfer(
        media_id    = item.id,
        kind        = "decommissioned",
        event_date  = date_type.today(),
        person_full_name = item.holder_full_name,
        department  = item.holder_department,
        operator    = current_user.username,
        notes       = reason,
    ))
    db.commit()
    db.refresh(item)
    return _media_to_out(item)


# ─── Выдача / возврат ───────────────────────────────────────────────────────

@router.post("/{item_id}/issue", response_model=MediaOut)
def issue_media(
        item_id:      int,
        payload:      IssueRequest,
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    item = db.query(MediaItem).filter(MediaItem.id == item_id).first()
    _check_unit(item, current_user)

    item.holder_full_name  = payload.holder_full_name
    item.holder_short_name = payload.holder_short_name
    item.holder_department = payload.holder_department
    item.issue_date        = payload.issue_date
    item.status            = "issued"

    db.add(MediaTransfer(
        media_id         = item.id,
        kind             = "issued",
        event_date       = payload.issue_date,
        person_full_name = payload.holder_full_name,
        department       = payload.holder_department,
        operator         = current_user.username,
        notes            = payload.notes,
    ))
    db.commit()
    db.refresh(item)
    return item


@router.post("/{item_id}/return", response_model=MediaOut)
def return_media(
        item_id:      int,
        return_date:  date_type = Query(..., description="Дата возврата"),
        notes:        Optional[str] = Query(None),
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    item = db.query(MediaItem).filter(MediaItem.id == item_id).first()
    _check_unit(item, current_user)

    # Запись в журнал — кому возвращаем (текущий держатель)
    db.add(MediaTransfer(
        media_id         = item.id,
        kind             = "returned",
        event_date       = return_date,
        person_full_name = item.holder_full_name,
        department       = item.holder_department,
        operator         = current_user.username,
        notes            = notes,
    ))

    item.holder_full_name  = None
    item.holder_short_name = None
    item.holder_department = None
    item.issue_date        = None
    item.status            = "available"

    db.commit()
    db.refresh(item)
    return item


@router.get("/{item_id}/transfers", response_model=list[TransferOut])
def list_transfers(
        item_id:      int,
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    item = db.query(MediaItem).filter(MediaItem.id == item_id).first()
    _check_unit(item, current_user)
    return [TransferOut.model_validate(t) for t in item.transfers]


# ─── Экспорт бирок в .docx ──────────────────────────────────────────────────

def _build_tags_docx(items: list[MediaItem]) -> io.BytesIO:
    """
    Сетка бирок: 5 колонок × N строк на A4-портрете. Каждая бирка ≈ 3.7×2.5 см
    (чтобы наклеить на флешку). Внутри: «{inv_number}», «{holder_short_name}»,
    «{issue_date dd.mm.yyyy}» — формат строго как в эталонной разметке.
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin   = Cm(1.0)
    section.right_margin  = Cm(1.0)

    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Бирки к машинным носителям информации")
    r.bold = True; r.font.size = Pt(11)

    # Только закреплённые носители — на свободные бирки не рисуем
    issued = [it for it in items if it.holder_full_name and it.inv_number]
    if not issued:
        doc.add_paragraph("Закреплённых носителей не найдено.").runs[0].font.size = Pt(10)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    COLS = 5
    rows = (len(issued) + COLS - 1) // COLS

    table = doc.add_table(rows=rows, cols=COLS)
    table.style = "Table Grid"
    table.autofit = False

    # Ширина каждой колонки ≈ 3.8 см (5×3.8 = 19, помещается в 19 см контента)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Cm(3.8)

    for idx, it in enumerate(issued):
        r_i = idx // COLS
        c_i = idx % COLS
        cell = table.rows[r_i].cells[c_i]

        # Очищаем дефолтный пустой параграф и пишем 3 строки
        cell.paragraphs[0].text = it.inv_number
        for r in cell.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p2.add_run(it.holder_short_name or it.holder_full_name or "")
        rn.font.size = Pt(9)

        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_str = it.issue_date.strftime("%d.%m.%Y") if it.issue_date else ""
        rd = p3.add_run(date_str)
        rd.font.size = Pt(9)

        # Все параграфы — без отступов
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(0)
            para.paragraph_format.line_spacing = 1.05

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── Импорт из Excel ────────────────────────────────────────────────────────
#
# Шаблон: 1-я строка — заголовки, со 2-й — данные. Маппинг русский header →
# поле модели (см. _IMPORT_COLUMNS). При импорте:
#   • Пустые строки и строки без инв. № пропускаются.
#   • Если инв. № уже есть в БД отдела — UPDATE (по согласию admin'а),
#     иначе INSERT.
#   • Поля типа/грифа/статуса принимают как русские лейблы, так и enum-коды.
#   • Даты: понимаем datetime/date из Excel + строки «ДД.ММ.ГГГГ» и «ГГГГ-ММ-ДД».
#   • Каждая строка с ошибкой — добавляется в errors[], не ломая остальной импорт.

# Упрощённый набор колонок: 6 полей. Гриф/статус/подразделение/полное ФИО
# подставляются автоматически — гриф всегда «ДСП», статус «Выдан» (если есть
# держатель), полное ФИО + подразделение — из общей базы людей по краткому.
# (excel_header, attr_name, required, example)
_IMPORT_COLUMNS = [
    ("Инв. №",             "inv_number",        True,  "1-ДСП"),
    ("Тип",                "media_type",        False, "Флешка USB"),
    ("Серийный №",         "serial_number",     False, "BCD-12345"),
    ("Объём, ГБ",          "capacity_gb",       False, 32),
    ("ФИО (краткое)",      "holder_short_name", False, "Иванов И.И."),
    ("Дата выдачи",        "issue_date",        False, "19.08.2024"),
    ("Примечание",         "notes",             False, ""),
]

# Карты «русское/произвольное → enum-код». Регистр игнорируется.
_TYPE_ALIASES = {
    "флешка": "flash", "флешка usb": "flash", "флеш": "flash", "usb": "flash",
    "ssd": "ssd", "внешний ssd": "ssd",
    "hdd": "hdd", "винчестер": "hdd", "внешний hdd": "hdd",
    "sd": "sd", "sd-карта": "sd", "карта памяти": "sd", "microsd": "sd",
    "cd": "cd_dvd", "dvd": "cd_dvd", "cd/dvd": "cd_dvd", "диск": "cd_dvd",
    "прочее": "other", "иное": "other",
}
_CLASS_ALIASES = {
    "открытый": "open", "о": "open",
    "дсп": "dsp", "для служебного пользования": "dsp",
    "секретно": "secret", "с": "secret",
    "совсекретно": "top_secret", "сов. секретно": "top_secret", "сс": "top_secret",
}
_STATUS_ALIASES = {
    "на хранении": "available", "свободен": "available", "в наличии": "available",
    "выдан": "issued",          "у пользователя": "issued",
    "неисправен": "broken",     "неработает": "broken", "сломан": "broken",
    "списан": "written_off",    "снят с учёта": "written_off",
    "утрачен": "lost",          "потерян": "lost",
}


def _normalize_enum(value: object, aliases: dict[str, str],
                    valid: tuple[str, ...]) -> Optional[str]:
    """Принимает русский лейбл, короткий код или None — возвращает enum-код."""
    if value is None or value == "":
        return None
    s = str(value).strip().lower()
    if s in valid:
        return s
    if s in aliases:
        return aliases[s]
    raise ValueError(f"Неизвестное значение «{value}». Допустимы: "
                     f"{', '.join(aliases.keys())}")


def _parse_date(value: object) -> Optional[date_type]:
    """Excel может прислать datetime или строку — нормализуем."""
    if value is None or value == "":
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date_type):
        return value
    s = str(value).strip()
    for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    raise ValueError(f"Не удалось распознать дату «{value}»")


class ImportRowError(BaseModel):
    row:     int
    message: str


class PersonCandidate(BaseModel):
    """Кандидат на сопоставление с краткой формой ФИО."""
    id:         int
    full_name:  str
    department: Optional[str] = None
    rank:       Optional[str] = None


class AmbiguousRow(BaseModel):
    """
    Сырая строка импорта, которая не однозначно сматчилась с базой людей —
    либо кандидатов несколько, либо ни одного. UI показывает оператору
    выбор и шлёт обратно через /import/resolve.
    """
    row:               int
    inv_number:        str
    media_type:        str
    serial_number:     Optional[str] = None
    capacity_gb:       Optional[int] = None
    holder_short_name: Optional[str] = None
    issue_date:        Optional[date_type] = None
    notes:             Optional[str] = None
    candidates:        list[PersonCandidate] = []


class ImportResult(BaseModel):
    added:     int
    updated:   int
    skipped:   int
    ambiguous: list[AmbiguousRow] = []
    errors:    list[ImportRowError] = []


class ResolvedRow(BaseModel):
    """Решение оператора по неоднозначной строке."""
    inv_number:        str
    media_type:        str = "flash"
    serial_number:     Optional[str] = None
    capacity_gb:       Optional[int] = None
    holder_short_name: Optional[str] = None
    issue_date:        Optional[date_type] = None
    notes:             Optional[str] = None
    person_id:         Optional[int] = None  # null = создать без полного ФИО


class ResolveRequest(BaseModel):
    rows: list[ResolvedRow]


# ─── Сопоставление краткого ФИО с базой людей ───────────────────────────────

def _parse_short_name(s: str) -> tuple[str, str]:
    """
    «Шевченко А.А.» → ("Шевченко", "АА")
    «Иванов И.» → ("Иванов", "И")
    «Иванов Иван Иванович» → ("Иванов", "ИИ")
    «Иванов» → ("Иванов", "")
    """
    s = (s or "").strip()
    if not s:
        return "", ""
    parts = s.split()
    surname = parts[0]
    rest = " ".join(parts[1:])
    # Убираем точки/пробелы/нижние подчёркивания, берём только инициалы
    initials = "".join(ch for ch in rest if ch.isalpha())
    # Если оператор написал «Иванов Иван Иванович» — взять только первые буквы слов
    if any(p.isalpha() and len(p) > 1 for p in parts[1:]):
        initials = "".join(p[0] for p in parts[1:] if p and p[0].isalpha())
    return surname, initials.upper()


def _match_persons_by_short(db: Session, short: str) -> list:
    """
    Ищет кандидатов в общей базе людей по краткому ФИО.
    Возвращает list[Person]. Учитывает только активных (fired_at IS NULL).
    """
    from app.models.person import Person  # внутри функции — избегаем cycle imports
    surname, initials = _parse_short_name(short)
    if not surname:
        return []
    query = db.query(Person).filter(
        Person.fired_at.is_(None),
        Person.full_name.ilike(f"{surname}%"),
    )
    candidates = query.limit(50).all()
    if not initials:
        return candidates

    # Фильтр по совпадению инициалов
    init_a = initials[0] if len(initials) >= 1 else ""
    init_b = initials[1] if len(initials) >= 2 else ""
    matched = []
    for p in candidates:
        words = (p.full_name or "").split()
        # Должна быть хотя бы фамилия + имя
        if len(words) < 2:
            continue
        # Сравниваем по первой букве имени и (опц.) отчества
        first_init = words[1][0].upper() if words[1] else ""
        if init_a and first_init != init_a:
            continue
        if init_b:
            patr_init = words[2][0].upper() if len(words) > 2 and words[2] else ""
            if patr_init != init_b:
                continue
        matched.append(p)
    return matched


def _persons_to_candidates(persons) -> list[PersonCandidate]:
    return [
        PersonCandidate(
            id=p.id, full_name=p.full_name,
            department=p.department, rank=p.rank,
        )
        for p in persons
    ]


@router.get("/import/template")
def download_import_template(
        count: int = Query(500, ge=1, le=1000,
                           description="Сколько строк с авто-нумерацией "
                                       "1-ДСП … N-ДСП пред-заполнить"),
        start: int = Query(1, ge=1, le=999_999,
                           description="С какого номера начать (по умолчанию 1)"),
        suffix: str = Query("ДСП", max_length=20,
                            description="Суффикс инв. номера (ДСП / С / СС / Открытый…)"),
        current_user: User = Depends(get_current_user),
):
    """
    Скачать .xlsx-шаблон с заголовками и пред-заполненной колонкой «Инв. №»
    (1-ДСП, 2-ДСП, … N-ДСП). Остальные колонки пустые — оператор только
    дописывает данные. Если строк мало, в Excel можно тянуть угол вниз —
    он сам продолжит нумерацию, узнав паттерн «{число}-{текст}».
    """
    if current_user.role not in ("admin", "unit"):
        raise HTTPException(status_code=403, detail="Доступ запрещён")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Учёт МНИ"

    HEADER_FILL  = PatternFill("solid", fgColor="4A7C6F")
    INV_FILL     = PatternFill("solid", fgColor="F5F4F2")  # лёгкая заливка колонки инв.№
    THIN   = Side(style="thin", color="CCCCCC")
    BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    WHITE  = Font(color="FFFFFF", bold=True, size=11)
    MONO   = Font(name="Cascadia Mono", size=11)

    # ── Заголовок ───────────────────────────────────────────────────────────
    for i, (label, _, required, _ex) in enumerate(_IMPORT_COLUMNS, start=1):
        c = ws.cell(row=1, column=i, value=label + (" *" if required else ""))
        c.fill = HEADER_FILL; c.font = WHITE; c.border = BORDER
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.column_dimensions[get_column_letter(i)].width = max(14, len(label) + 2)
    ws.row_dimensions[1].height = 36

    # Заморозим шапку
    ws.freeze_panes = "A2"

    # ── Пред-заполнение колонки «Инв. №» ────────────────────────────────────
    # _IMPORT_COLUMNS[0] всегда «Инв. №» — поэтому колонка 1.
    suffix_clean = (suffix or "ДСП").strip()
    for i in range(count):
        row_idx = i + 2  # строки начинаются со 2-й (1-я — заголовок)
        n       = start + i
        cell = ws.cell(row=row_idx, column=1, value=f"{n}-{suffix_clean}")
        cell.font   = MONO
        cell.fill   = INV_FILL
        cell.border = BORDER
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # ── Подсказка по допустимым значениям (под таблицей) ────────────────────
    note_start = count + 4
    notes = [
        "Поле со * — обязательное. Колонка «Инв. №» уже пред-заполнена — оставьте",
        "номер, удалите или отредактируйте по необходимости.",
        "",
        "Допустимые значения:",
        "  Тип:    Флешка USB · SSD · HDD · SD-карта · CD/DVD · Прочее",
        "  Гриф:   Открытый · ДСП · Секретно · Совсекретно",
        "  Статус: На хранении · Выдан · Неисправен · Списан · Утрачен",
        "  Даты:   «19.08.2024» или формат даты Excel",
    ]
    bold = Font(bold=True, color="666666")
    plain = Font(color="888888")
    for offset, text in enumerate(notes):
        c = ws.cell(row=note_start + offset, column=1, value=text)
        c.font = bold if offset in (0, 3) else plain

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"Учёт_МНИ_шаблон_{count}_строк.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition":
                f'attachment; filename*=UTF-8\'\'{quote(filename)}',
        },
    )


def _parse_excel(content: bytes) -> tuple[list[dict], list[ImportRowError]]:
    """
    Парсит .xlsx по нашему упрощённому шаблону. Возвращает чистые словари
    и список ошибок парсинга (типы/даты/числа). Сами решения о записи
    в БД не принимает — этим занимается /import (preview) и /resolve.
    """
    try:
        wb = openpyxl.load_workbook(io.BytesIO(content),
                                    data_only=True, read_only=True)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Не удалось прочитать .xlsx: {e}")

    ws = wb.active
    rows = ws.iter_rows(values_only=True)

    header_row = next(rows, None)
    if not header_row:
        raise HTTPException(status_code=400, detail="Файл пуст")

    header_map: dict[int, str] = {}
    excel_label_to_attr = {label.lower(): attr for label, attr, *_ in _IMPORT_COLUMNS}
    for col_idx, cell_val in enumerate(header_row):
        if cell_val is None:
            continue
        label = str(cell_val).replace("*", "").strip().lower()
        if label in excel_label_to_attr:
            header_map[col_idx] = excel_label_to_attr[label]

    if "inv_number" not in header_map.values():
        raise HTTPException(
            status_code=400,
            detail="В файле не найдена колонка «Инв. №». "
                   "Скачайте свежий шаблон через «📥 Шаблон».",
        )

    parsed: list[dict] = []
    errors: list[ImportRowError] = []

    for row_idx, row in enumerate(rows, start=2):
        if not any(row):
            continue
        data: dict[str, object] = {"_row": row_idx}
        for col_idx, val in enumerate(row):
            attr = header_map.get(col_idx)
            if not attr:
                continue
            if isinstance(val, str):
                val = val.strip()
                if val == "":
                    val = None
            data[attr] = val

        if not data.get("inv_number"):
            continue   # пустые строки — тихо пропускаем

        # Нормализация
        try:
            data["media_type"] = _normalize_enum(
                data.get("media_type"), _TYPE_ALIASES, MEDIA_TYPES
            ) or "flash"
            data["issue_date"] = _parse_date(data.get("issue_date"))
        except ValueError as ve:
            errors.append(ImportRowError(row=row_idx, message=str(ve)))
            continue

        cap = data.get("capacity_gb")
        if cap is not None and cap != "":
            try:
                # Принимаем «64 Гб», «64гб», «64» — извлекаем число
                if isinstance(cap, str):
                    digits = "".join(ch for ch in cap if ch.isdigit() or ch == ".")
                    cap = float(digits) if digits else 0
                data["capacity_gb"] = int(float(cap))
            except (ValueError, TypeError):
                errors.append(ImportRowError(
                    row=row_idx,
                    message=f"Неверный объём «{cap}» — ожидается число"))
                continue

        data["inv_number"] = str(data["inv_number"]).strip()
        parsed.append(data)

    return parsed, errors


def _detect_duplicate_serials(
        db: Session, unit_name: str, parsed: list[dict],
) -> tuple[set[int], list[ImportRowError]]:
    """
    Проверяет серийные номера на дубли:
      • в самом файле — если две строки с одним серийником, обе помечаются ошибкой;
      • в БД — если в файле есть серийник, который уже есть на другом носителе
        (другой инв.№), строка тоже помечается ошибкой.
    Возвращает (set индексов строк-проблем, список ошибок).
    Серийник «-» / пустой / '0' — не считается значимым.
    """
    bad_rows: set[int] = set()
    errors: list[ImportRowError] = []

    def _is_meaningful(s: object) -> bool:
        if not s:
            return False
        s = str(s).strip()
        return s not in ("", "-", "—", "0", "0000", "n/a", "N/A")

    # 1. Дубли внутри файла
    serial_to_rows: dict[str, list[dict]] = {}
    for d in parsed:
        sn = d.get("serial_number")
        if not _is_meaningful(sn):
            continue
        serial_to_rows.setdefault(str(sn).strip(), []).append(d)

    for sn, dupes in serial_to_rows.items():
        if len(dupes) < 2:
            continue
        for d in dupes:
            row_idx = d.get("_row", 0)
            bad_rows.add(row_idx)
            other_invs = ", ".join(
                f"«{x['inv_number']}» (стр. {x.get('_row')})"
                for x in dupes if x is not d
            )
            errors.append(ImportRowError(
                row=row_idx,
                message=f"Серийный «{sn}» встречается в файле повторно — ещё в {other_invs}",
            ))

    # 2. Дубли с БД (другой инв.№, тот же отдел)
    file_serials = {
        str(d.get("serial_number")).strip()
        for d in parsed if _is_meaningful(d.get("serial_number"))
    }
    if file_serials:
        existing = (
            db.query(MediaItem.serial_number, MediaItem.inv_number)
              .filter(
                  MediaItem.unit_username == unit_name,
                  MediaItem.serial_number.in_(file_serials),
              )
              .all()
        )
        existing_map = {sn: inv for sn, inv in existing}
        for d in parsed:
            sn = d.get("serial_number")
            if not _is_meaningful(sn):
                continue
            sn = str(sn).strip()
            inv_in_file = d["inv_number"]
            inv_in_db = existing_map.get(sn)
            if inv_in_db and inv_in_db != inv_in_file:
                row_idx = d.get("_row", 0)
                bad_rows.add(row_idx)
                errors.append(ImportRowError(
                    row=row_idx,
                    message=f"Серийный «{sn}» уже закреплён за носителем «{inv_in_db}» в БД",
                ))

    return bad_rows, errors


def _save_resolved_row(
        db: Session, unit_name: str,
        data: dict, person_id: Optional[int],
) -> str:
    """
    Сохраняет/обновляет MediaItem по разрешённой строке.
    Если person_id задан — подтягивает full_name + department из персоны.
    Возвращает 'added' / 'updated'.
    """
    full_name: Optional[str] = None
    department: Optional[str] = None
    if person_id:
        from app.models.person import Person
        p = db.query(Person).filter(Person.id == person_id).first()
        if p:
            full_name  = p.full_name
            department = p.department

    inv = data["inv_number"]
    short = data.get("holder_short_name") or None
    has_holder = bool(full_name or short)

    payload_attrs = {
        "media_type":        data.get("media_type") or "flash",
        "serial_number":     data.get("serial_number"),
        "capacity_gb":       data.get("capacity_gb"),
        "classification":    "dsp",                # всегда ДСП
        "holder_person_id":  person_id,            # связь с общей базой
        "holder_full_name":  full_name,
        "holder_short_name": short,
        "holder_department": department,
        "issue_date":        data.get("issue_date"),
        # дата выдачи = последняя проверка (по требованию)
        "last_check_date":   data.get("issue_date"),
        "notes":             data.get("notes"),
        "status":            "issued" if has_holder else "available",
    }

    existing = (
        db.query(MediaItem)
          .filter(MediaItem.unit_username == unit_name,
                  MediaItem.inv_number    == inv)
          .first()
    )
    if existing:
        for k, v in payload_attrs.items():
            if v is not None or k in ("notes",):  # notes допускаем None
                setattr(existing, k, v)
        return "updated"
    else:
        item = MediaItem(unit_username=unit_name, inv_number=inv, **payload_attrs)
        db.add(item)
        return "added"


@router.post("/import", response_model=ImportResult)
async def import_media(
        file:         UploadFile = File(...),
        unit:         Optional[str] = Query(None, max_length=100),
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    """
    Парсит .xlsx и пытается сматчить краткое ФИО с общей базой людей.
      • 1 кандидат  → строка сохраняется сразу (full_name + department берутся из персоны)
      • 0 или >1    → возвращается в `ambiguous[]` для разрешения оператором

    Если ФИО в файле нет — носитель сохраняется без держателя (status=available).
    """
    unit_name = _resolve_unit(current_user, unit)

    if not file.filename or not file.filename.lower().endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Ожидается файл .xlsx")
    contents = await file.read()
    if len(contents) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Файл больше 10 МБ")

    parsed, errors = _parse_excel(contents)

    # Детект дублей серийных номеров (в файле и против БД).
    # Грязные строки в bad_rows — пропускаем, чтобы не задвоить.
    bad_rows, dup_errors = _detect_duplicate_serials(db, unit_name, parsed)
    errors.extend(dup_errors)

    added = 0; updated = 0; skipped = 0
    ambiguous: list[AmbiguousRow] = []

    for data in parsed:
        row_idx = data.get("_row", 0)
        if row_idx in bad_rows:
            skipped += 1
            continue
        try:
            short = data.get("holder_short_name")
            person_id: Optional[int] = None

            if short:
                matches = _match_persons_by_short(db, short)
                if len(matches) == 1:
                    person_id = matches[0].id
                elif len(matches) == 0:
                    # Никого не нашли — возвращаем для resolve, без кандидатов
                    ambiguous.append(AmbiguousRow(
                        row=row_idx,
                        inv_number=data["inv_number"],
                        media_type=data.get("media_type") or "flash",
                        serial_number=data.get("serial_number"),
                        capacity_gb=data.get("capacity_gb"),
                        holder_short_name=short,
                        issue_date=data.get("issue_date"),
                        notes=data.get("notes"),
                        candidates=[],
                    ))
                    continue
                else:
                    ambiguous.append(AmbiguousRow(
                        row=row_idx,
                        inv_number=data["inv_number"],
                        media_type=data.get("media_type") or "flash",
                        serial_number=data.get("serial_number"),
                        capacity_gb=data.get("capacity_gb"),
                        holder_short_name=short,
                        issue_date=data.get("issue_date"),
                        notes=data.get("notes"),
                        candidates=_persons_to_candidates(matches),
                    ))
                    continue

            kind = _save_resolved_row(db, unit_name, data, person_id)
            if   kind == "added":   added   += 1
            elif kind == "updated": updated += 1
            db.commit()
        except Exception as e:
            db.rollback()
            errors.append(ImportRowError(row=row_idx, message=str(e)))

    return ImportResult(
        added=added, updated=updated, skipped=skipped,
        ambiguous=ambiguous, errors=errors,
    )


@router.post("/import/resolve", response_model=ImportResult)
def resolve_import(
        payload:      ResolveRequest,
        unit:         Optional[str] = Query(None, max_length=100),
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    """
    Дописывает строки, по которым оператор сделал выбор в UI.
    person_id=null → сохранить без полного ФИО (только краткое).
    """
    unit_name = _resolve_unit(current_user, unit)
    added = 0; updated = 0
    errors: list[ImportRowError] = []
    for i, row in enumerate(payload.rows, start=1):
        try:
            data = row.model_dump()
            kind = _save_resolved_row(db, unit_name, data, row.person_id)
            if   kind == "added":   added   += 1
            elif kind == "updated": updated += 1
            db.commit()
        except Exception as e:
            db.rollback()
            errors.append(ImportRowError(row=i, message=str(e)))
    return ImportResult(added=added, updated=updated, skipped=0,
                         ambiguous=[], errors=errors)


@router.get("/tags-export")
def export_tags(
        unit:         Optional[str] = Query(None, max_length=100),
        db:           Session = Depends(get_db),
        current_user: User    = Depends(get_current_user),
):
    unit_name = _resolve_unit(current_user, unit)
    items = (
        db.query(MediaItem)
          .filter(MediaItem.unit_username == unit_name)
          .order_by(MediaItem.inv_number.asc())
          .all()
    )
    buf = _build_tags_docx(items)
    filename = f"Бирки_МНИ_{unit_name}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition":
                f'attachment; filename*=UTF-8\'\'{quote(filename)}',
        },
    )
