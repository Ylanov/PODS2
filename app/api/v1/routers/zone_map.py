# app/api/v1/routers/zone_map.py
"""
Карта зон — pods2-сторона (вкладка под permission='zone_map').

Самостоятельный дубль карты Оперативного дежурного, заточенный под импорт
координат из Excel. CRUD зон + разбор .xlsx + шаблон. Тайлы карты и сам
Leaflet берём из ПУБЛИЧНЫХ прокси карты ОД (/api/v1/oper-map/tile/...,
/api/v1/oper-map/vendor/...) — они уже проксируют Яндекс/CDN через
интернет-сетевуху сервера, дублировать их незачем.

Эндпоинты (все под permission='zone_map'; admin проходит всегда):

  GET    /api/v1/zone-map/zones            — все зоны
  POST   /api/v1/zone-map/zones            — новая зона
  PATCH  /api/v1/zone-map/zones/{id}
  DELETE /api/v1/zone-map/zones/{id}
  DELETE /api/v1/zone-map/zones            — очистить все зоны
  POST   /api/v1/zone-map/import           — разбор Excel (multipart) → зоны
  GET    /api/v1/zone-map/template.xlsx    — шаблон для заполнения
"""

import io
import logging
import re
from typing import List, Optional

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, Form, HTTPException, Response, UploadFile
from openpyxl import Workbook, load_workbook
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from app.api.dependencies import require_permission
from app.core.geodesy import msk77_to_wgs84
from app.db.database import get_db
from app.models.zone_map import ZoneMapZone


logger = logging.getLogger(__name__)

router = APIRouter(dependencies=[Depends(require_permission("zone_map"))])


# Палитра для авто-раскраски зон, если в Excel нет колонки «Цвет».
# Различимые насыщенные цвета — чтобы соседние зоны не сливались.
_PALETTE = [
    "#1976d2", "#e53935", "#43a047", "#fb8c00", "#8e24aa",
    "#00897b", "#c2185b", "#3949ab", "#7cb342", "#f4511e",
    "#00acc1", "#6d4c41", "#d81b60", "#5e35b1", "#039be5",
]

_HEX_RE = re.compile(r"^#(?:[0-9a-fA-F]{3}|[0-9a-fA-F]{6})$")


# ─── Pydantic ────────────────────────────────────────────────────────────────

class ZoneOut(BaseModel):
    id:         int
    name:       str
    role:       Optional[str] = None
    color:      str
    points:     list   # [[lat, lng], ...]
    sort_order: int


class ZoneIn(BaseModel):
    name:       str = Field(..., min_length=1, max_length=200)
    role:       Optional[str] = Field(default=None, max_length=200)
    color:      str = Field(default="#1976d2", max_length=20)
    points:     list = Field(default_factory=list)
    sort_order: int = 0


class ZonePatch(BaseModel):
    name:       Optional[str] = Field(default=None, min_length=1, max_length=200)
    role:       Optional[str] = Field(default=None, max_length=200)
    color:      Optional[str] = Field(default=None, max_length=20)
    points:     Optional[list] = None
    sort_order: Optional[int] = None


class ImportResult(BaseModel):
    zones_created: int
    points_total:  int
    rows_skipped:  int
    sheet:         Optional[str] = None
    columns:       dict          # как распознаны колонки (для подсказки в UI)
    zones:         List[ZoneOut]


# ─── helpers ──────────────────────────────────────────────────────────────────

def _zone_out(z: ZoneMapZone) -> ZoneOut:
    return ZoneOut(
        id=z.id,
        name=z.name,
        role=z.role,
        color=z.color,
        points=z.get_points(),
        sort_order=z.sort_order,
    )


def _clean_points(raw: list) -> list:
    """Оставляет только валидные [lat, lng] пары в допустимых диапазонах."""
    out = []
    for p in raw or []:
        try:
            lat = float(p[0])
            lng = float(p[1])
        except (TypeError, ValueError, IndexError):
            continue
        if -90.0 <= lat <= 90.0 and -180.0 <= lng <= 180.0:
            out.append([lat, lng])
    return out


# ─── CRUD ─────────────────────────────────────────────────────────────────────

@router.get("/zones", response_model=List[ZoneOut], summary="Список зон")
def list_zones(db: Session = Depends(get_db)):
    rows = (
        db.query(ZoneMapZone)
        .order_by(ZoneMapZone.sort_order.asc(), ZoneMapZone.id.asc())
        .all()
    )
    return [_zone_out(z) for z in rows]


@router.post("/zones", response_model=ZoneOut, status_code=201, summary="Создать зону")
def create_zone(payload: ZoneIn, db: Session = Depends(get_db)):
    z = ZoneMapZone(
        name=payload.name.strip(),
        role=(payload.role or "").strip() or None,
        color=payload.color or "#1976d2",
        sort_order=payload.sort_order or 0,
    )
    z.set_points(_clean_points(payload.points))
    db.add(z)
    db.commit()
    db.refresh(z)
    return _zone_out(z)


@router.patch("/zones/{zone_id}", response_model=ZoneOut, summary="Обновить зону")
def patch_zone(zone_id: int, payload: ZonePatch, db: Session = Depends(get_db)):
    z = db.query(ZoneMapZone).filter(ZoneMapZone.id == zone_id).first()
    if not z:
        raise HTTPException(status_code=404, detail="Зона не найдена")
    if payload.name is not None:
        z.name = payload.name.strip()
    if payload.role is not None:
        z.role = payload.role.strip() or None
    if payload.color is not None:
        z.color = payload.color
    if payload.points is not None:
        z.set_points(_clean_points(payload.points))
    if payload.sort_order is not None:
        z.sort_order = payload.sort_order
    db.commit()
    db.refresh(z)
    return _zone_out(z)


@router.delete("/zones/{zone_id}", status_code=204, summary="Удалить зону")
def delete_zone(zone_id: int, db: Session = Depends(get_db)):
    z = db.query(ZoneMapZone).filter(ZoneMapZone.id == zone_id).first()
    if not z:
        raise HTTPException(status_code=404, detail="Зона не найдена")
    db.delete(z)
    db.commit()


@router.delete("/zones", status_code=204, summary="Очистить все зоны")
def clear_zones(db: Session = Depends(get_db)):
    db.query(ZoneMapZone).delete()
    db.commit()


# ─── Разбор Excel ─────────────────────────────────────────────────────────────
#
# Колонки распознаём по заголовку (любой регистр, синонимы). Минимально нужны
# широта и долгота; колонка «зона» группирует точки в полигоны. Если зоны нет —
# все точки идут в одну зону. Если заголовков нет вовсе — берём позиционно:
# [зона, широта, долгота] или [широта, долгота].

_LAT_KEYS  = ("широта", "шир", "lat", "latitude")
_LNG_KEYS  = ("долгота", "долг", "lng", "lon", "long", "longitude")
# Для режима МСК-77: колонки прямоугольных координат (X — север, Y — восток).
_X_KEYS    = ("x", "север", "north", "абсцисса")
_Y_KEYS    = ("y", "восток", "east", "ордината")
_ZONE_KEYS = ("зона", "зоны", "участок", "группа", "group", "zone",
              "наименование", "название", "объект", "район", "name")
_ROLE_KEYS = ("роль", "категория", "тип", "role", "подпись", "примечание", "комментарий")
_COLOR_KEYS = ("цвет", "color", "colour")
_COMBO_KEYS = ("координаты", "коорд", "coords", "coordinates", "latlng", "ll")


def _norm_header(v) -> str:
    return re.sub(r"[^a-zа-я0-9]", "", str(v or "").strip().lower())


def _match_col(header: str, keys) -> bool:
    h = _norm_header(header)
    if not h:
        return False
    return any(_norm_header(k) in h for k in keys)


def _to_float(v) -> Optional[float]:
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace(",", ".")
    # «55.7558°» / «55.7558 N» — оставляем только число со знаком
    m = re.search(r"-?\d+(?:\.\d+)?", s)
    if not m:
        return None
    try:
        return float(m.group(0))
    except ValueError:
        return None


def _split_combo(v):
    """Разбирает «55.75, 37.61» (или ; / пробел) в пару чисел."""
    if v is None:
        return None
    parts = re.split(r"[;,\s]+", str(v).strip())
    nums = [_to_float(p) for p in parts if p]
    nums = [n for n in nums if n is not None]
    if len(nums) >= 2:
        return nums[0], nums[1]
    return None


def _orient(a: float, b: float):
    """
    Возвращает (lat, lng). Широта всегда |≤90|, долгота может быть больше.
    Если a по модулю >90 — значит это долгота, меняем местами. Иначе считаем,
    что порядок (широта, долгота) — как в шаблоне.
    """
    if abs(a) > 90 >= abs(b):
        return b, a
    return a, b


def _extract_rows_xlsx(data: bytes):
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001 — любой сбой парсинга = понятная 400
        raise HTTPException(status_code=400, detail=f"Не удалось прочитать Excel: {exc}")
    ws = wb.active
    rows = [list(r) for r in ws.iter_rows(values_only=True)]
    rows = [r for r in rows if any(c is not None and str(c).strip() != "" for c in r)]
    return rows, ws.title


def _extract_rows_docx(data: bytes):
    """Берёт первую таблицу из .docx и возвращает её строки как список списков."""
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=400, detail=f"Не удалось прочитать Word: {exc}")
    if not doc.tables:
        raise HTTPException(status_code=400, detail="В документе нет таблицы с координатами")
    tb = doc.tables[0]
    rows = []
    for row in tb.rows:
        cells = [c.text.strip() for c in row.cells]
        if any(cells):
            rows.append(cells)
    return rows, "Word"


def _cell(row, idx):
    return row[idx] if idx is not None and idx < len(row) else None


def _detect_columns(rows, coord_system):
    """
    Ищет строку заголовка среди первых 6 строк и индексы колонок.
    Для wgs84 — широта/долгота (или объединённая колонка координат); для msk77 — X/Y.
    Дополнительно ищет зону/подпись/цвет. Возвращает (cols, data_rows).
    """
    is_msk = coord_system == "msk77"
    found = {"lat": None, "lng": None, "x": None, "y": None,
             "zone": None, "role": None, "color": None, "combo": None}
    header_row = -1
    for i, row in enumerate(rows[:6]):
        cand = dict.fromkeys(found, None)
        for ci, cell in enumerate(row):
            if is_msk:
                if cand["x"] is None and _match_col(cell, _X_KEYS):
                    cand["x"] = ci
                elif cand["y"] is None and _match_col(cell, _Y_KEYS):
                    cand["y"] = ci
            else:
                if cand["lat"] is None and _match_col(cell, _LAT_KEYS):
                    cand["lat"] = ci
                elif cand["lng"] is None and _match_col(cell, _LNG_KEYS):
                    cand["lng"] = ci
                elif cand["combo"] is None and _match_col(cell, _COMBO_KEYS):
                    cand["combo"] = ci
            if cand["zone"] is None and _match_col(cell, _ZONE_KEYS):
                cand["zone"] = ci
            if cand["role"] is None and _match_col(cell, _ROLE_KEYS):
                cand["role"] = ci
            if cand["color"] is None and _match_col(cell, _COLOR_KEYS):
                cand["color"] = ci
        ok = (cand["x"] is not None and cand["y"] is not None) if is_msk else \
             ((cand["lat"] is not None and cand["lng"] is not None) or cand["combo"] is not None)
        if ok:
            found = cand
            header_row = i
            break

    if header_row >= 0:
        data_rows = rows[header_row + 1:]
    else:
        # Заголовков не нашли — позиционно. Считаем, что есть ≥2 числовых колонки.
        ncols = max((len(r) for r in rows), default=0)
        first_nums = [c for c in (rows[0] if rows else []) if _to_float(c) is not None]
        data_rows = rows if len(first_nums) >= 2 else rows[1:]
        if ncols >= 3:
            found["zone"] = 0
            if is_msk:
                found["x"], found["y"] = 1, 2
            else:
                found["lat"], found["lng"] = 1, 2
        else:
            if is_msk:
                found["x"], found["y"] = 0, 1
            else:
                found["lat"], found["lng"] = 0, 1

    found["header_row"] = header_row + 1 if header_row >= 0 else None
    return found, data_rows


def _rows_to_zones(rows, coord_system, default_name):
    rows = [r for r in rows if any(c is not None and str(c).strip() != "" for c in r)]
    if not rows:
        raise HTTPException(status_code=400, detail="Файл пустой")
    is_msk = coord_system == "msk77"
    cols, data_rows = _detect_columns(rows, coord_system)

    groups: dict = {}
    order: list = []
    skipped = 0

    for row in data_rows:
        lat = lng = None
        if is_msk:
            x = _to_float(_cell(row, cols["x"]))
            y = _to_float(_cell(row, cols["y"]))
            # Отсекаем мусор/строку нумерации колонок («1,2,3»): настоящая
            # абсцисса МСК-77 для Москвы — тысячи метров, не единицы.
            if x is None or y is None or abs(x) < 100:
                skipped += 1
                continue
            try:
                lat, lng = msk77_to_wgs84(x, y)
            except Exception:  # noqa: BLE001
                skipped += 1
                continue
        else:
            if cols["combo"] is not None:
                pair = _split_combo(_cell(row, cols["combo"]))
                if pair:
                    lat, lng = _orient(pair[0], pair[1])
            if lat is None or lng is None:
                a = _to_float(_cell(row, cols["lat"]))
                b = _to_float(_cell(row, cols["lng"]))
                if a is None or b is None:
                    skipped += 1
                    continue
                lat, lng = _orient(a, b)

        if not (-90.0 <= lat <= 90.0 and -180.0 <= lng <= 180.0):
            skipped += 1
            continue

        zname = _cell(row, cols["zone"])
        zname = str(zname).strip() if zname is not None and str(zname).strip() else default_name
        if zname not in groups:
            groups[zname] = {"points": [], "color": None, "role": None}
            order.append(zname)
        groups[zname]["points"].append([round(lat, 7), round(lng, 7)])

        cval = _cell(row, cols["color"])
        if cval and groups[zname]["color"] is None:
            cs = str(cval).strip()
            cs = cs if cs.startswith("#") else f"#{cs}"
            if _HEX_RE.match(cs):
                groups[zname]["color"] = cs
        rval = _cell(row, cols["role"])
        if rval and groups[zname]["role"] is None:
            groups[zname]["role"] = str(rval).strip()

    if not order:
        hint = ("колонки X и Y (МСК-77)" if is_msk
                else "колонки с широтой и долготой")
        raise HTTPException(
            status_code=400,
            detail=f"Не нашли ни одной валидной координаты. Проверьте, что есть {hint} "
                   "(или скачайте шаблон).",
        )

    parsed = []
    for i, name in enumerate(order):
        g = groups[name]
        parsed.append({
            "name":  name,
            "role":  g["role"],
            "color": g["color"] or _PALETTE[i % len(_PALETTE)],
            "points": g["points"],
        })
    return parsed, skipped, cols


def _filename_stem(filename: str) -> str:
    base = (filename or "").rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    stem = base.rsplit(".", 1)[0].strip()
    return stem[:200] or "Импорт"


@router.post("/import", response_model=ImportResult, summary="Импорт зон из Excel/Word")
async def import_zones(
    file: UploadFile = File(...),
    mode: str = Form("replace"),                 # replace | append
    coord_system: str = Form("wgs84"),           # wgs84 | msk77
    db: Session = Depends(get_db),
):
    name = (file.filename or "").lower()
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Файл пустой")

    if name.endswith((".xlsx", ".xlsm")):
        rows, sheet = _extract_rows_xlsx(data)
    elif name.endswith(".docx"):
        rows, sheet = _extract_rows_docx(data)
    else:
        raise HTTPException(status_code=400, detail="Поддерживаются файлы .xlsx и .docx")

    default_name = _filename_stem(file.filename)
    parsed, skipped, columns_info = _rows_to_zones(rows, coord_system, default_name)

    if mode == "replace":
        db.query(ZoneMapZone).delete()
        base_sort = 0
    else:
        base_sort = (db.query(ZoneMapZone).count())

    created = []
    points_total = 0
    for i, p in enumerate(parsed):
        z = ZoneMapZone(
            name=p["name"][:200],
            role=(p["role"] or None),
            color=p["color"],
            sort_order=base_sort + i,
        )
        z.set_points(p["points"])
        db.add(z)
        created.append(z)
        points_total += len(p["points"])
    db.commit()
    for z in created:
        db.refresh(z)

    return ImportResult(
        zones_created=len(created),
        points_total=points_total,
        rows_skipped=skipped,
        sheet=sheet,
        columns=columns_info,
        zones=[_zone_out(z) for z in created],
    )


# ─── Шаблон Excel ─────────────────────────────────────────────────────────────

@router.get("/template.xlsx", summary="Скачать шаблон Excel для импорта зон")
def download_template():
    wb = Workbook()
    ws = wb.active
    ws.title = "Зоны"
    ws.append(["Зона", "Широта", "Долгота", "Цвет (необяз.)", "Подпись (необяз.)"])
    # Пример: две зоны по несколько вершин (Москва, около центра).
    sample = [
        ["Зона 1", 55.7558, 37.6173, "#e53935", "Штаб"],
        ["Зона 1", 55.7570, 37.6210, "", ""],
        ["Зона 1", 55.7540, 37.6225, "", ""],
        ["Зона 1", 55.7535, 37.6175, "", ""],
        ["Зона 2", 55.7600, 37.6100, "#43a047", "Объект А"],
        ["Зона 2", 55.7615, 37.6140, "", ""],
        ["Зона 2", 55.7595, 37.6160, "", ""],
    ]
    for r in sample:
        ws.append(r)
    # Чуть шире колонки, чтобы заголовки читались.
    for col, width in zip("ABCDE", (16, 12, 12, 16, 18)):
        ws.column_dimensions[col].width = width

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="zone_map_template.xlsx"'},
    )
