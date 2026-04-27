# app/api/v1/routers/comms_report.py
"""
API для отчёта Форма 3-СВЯЗЬ (отдел связи).

Эндпоинты:
  GET  /comms-report?year=YYYY            — получить отчёт за год
                                            (создаётся с дефолтами при первом обращении)
  PUT  /comms-report?year=YYYY            — сохранить весь снимок (data)
  POST /comms-report/export?year=YYYY     — выгрузить .docx, готовый к подписи

Доступ:
  • admin           — видит/правит отчёт любого отдела (указывает ?unit=<username>,
                      иначе — свой логин)
  • role='unit'     — видит/правит ТОЛЬКО свой отчёт (unit_username = username)
  • остальные роли  — 403
"""

import io
from datetime import datetime, timezone
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, ConfigDict, Field
from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from app.db.database import get_db
from app.models.user import User
from app.models.comms_report import CommsReport
from app.api.dependencies import get_current_user
from app.data.comms_report_defaults import (
    make_default_report, compute_derived, normalize_item,
)


router = APIRouter()


# ─── Схемы ──────────────────────────────────────────────────────────────────

class CommsReportItem(BaseModel):
    """Одна позиция внутри направления.
    Поля 12-19 docx (nz / td / backup_fund / mchs_reserve / capital_repair /
    mb / written_off / plus) — это «кроме того» подкатегории."""
    id:             str
    name:           str
    is_group:       bool = False
    required:       int  = 0
    start:          int  = 0
    arrived:        int  = 0
    removed:        int  = 0
    working:        int  = 0
    modern:         int  = 0
    overdue:        int  = 0
    nz:             int  = 0   # В «НЗ»
    td:             int  = 0   # На «ТД»
    backup_fund:    int  = 0   # Подменный фонд 2-3 кат
    mchs_reserve:   int  = 0   # Резерв МЧС России
    capital_repair: int  = 0   # В т.ч. кап. ремонт
    mb:             int  = 0   # На МБ
    written_off:    int  = 0   # В т.ч. списано
    plus:           int  = 0   # В запасах центров
    note:           str  = ""


class CommsReportCategory(BaseModel):
    id:    str
    index: int
    title: str
    unit:  str
    items: list[CommsReportItem]


class CommsReportResponse(BaseModel):
    unit_username: str
    year:          int
    data:          list[CommsReportCategory]
    updated_at:    datetime

    model_config = ConfigDict(from_attributes=True)


class CommsReportUpdate(BaseModel):
    data: list[CommsReportCategory] = Field(default_factory=list)


# ─── Вспомогательные функции ────────────────────────────────────────────────

def _resolve_unit(current_user: User, unit_override: str | None) -> str:
    """
    Определяет username отдела, чей отчёт запрашивается.
    Admin может указать ?unit=xxx, отделу-пользователю всегда подставляется
    его собственный username.
    """
    if current_user.role == "admin":
        return (unit_override or current_user.username).strip()
    if current_user.role == "unit":
        # unit-юзер может смотреть только свой отчёт — override игнорируем.
        return current_user.username
    raise HTTPException(status_code=403, detail="Доступ только для отдела или админа")


def _get_or_create_report(db: Session, unit: str, year: int) -> CommsReport:
    """Получает отчёт из БД или создаёт новый с дефолтной структурой."""
    rep = (
        db.query(CommsReport)
          .filter(CommsReport.unit_username == unit, CommsReport.year == year)
          .first()
    )
    if rep:
        return rep

    rep = CommsReport(
        unit_username = unit,
        year          = year,
        data          = make_default_report(),
    )
    db.add(rep)
    db.commit()
    db.refresh(rep)
    return rep


# ─── GET: получить отчёт ────────────────────────────────────────────────────

@router.get("", response_model=CommsReportResponse)
def get_report(
        year:         int              = Query(..., ge=2020, le=2100),
        unit:         str | None       = Query(None, max_length=100),
        db:           Session          = Depends(get_db),
        current_user: User             = Depends(get_current_user),
):
    unit_name = _resolve_unit(current_user, unit)
    rep = _get_or_create_report(db, unit_name, year)
    return rep


# ─── PUT: сохранить отчёт ───────────────────────────────────────────────────

@router.put("", response_model=CommsReportResponse)
def update_report(
        payload:      CommsReportUpdate,
        year:         int              = Query(..., ge=2020, le=2100),
        unit:         str | None       = Query(None, max_length=100),
        db:           Session          = Depends(get_db),
        current_user: User             = Depends(get_current_user),
):
    unit_name = _resolve_unit(current_user, unit)
    rep = _get_or_create_report(db, unit_name, year)

    # Нормализуем числа и сохраняем как есть (без пересчёта производных —
    # они вычисляются на клиенте; но для защиты приводим типы).
    data = []
    for cat in payload.data:
        data.append({
            "id":    cat.id,
            "index": cat.index,
            "title": cat.title,
            "unit":  cat.unit,
            "items": [normalize_item(it.model_dump()) for it in cat.items],
        })
    rep.data = data
    # SQLAlchemy не замечает изменения внутри JSONB без flag_modified,
    # но здесь присваивание целого списка — replacement, изменение видно.
    db.commit()
    db.refresh(rep)
    return rep


# ─── POST: экспорт в .docx ──────────────────────────────────────────────────

# Институциональный бланк. Все строки идут БЕЗ пустых абзацев между ними —
# в оригинале расстояния минимальные. Если в визуале хочется отступ —
# увеличиваем space_after у конкретного абзаца, не пустыми параграфами.
_HEADER_LEFT = [
    ("МЧС РОССИИ",                                                   {"bold": True, "size": 10, "align": "center"}),
    ("ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ",                                  {"bold": True, "size": 8,  "align": "center", "space_before": True}),
    ("КАЗЕННОЕ УЧРЕЖДЕНИЕ",                                          {"bold": True, "size": 8,  "align": "center"}),
    ("«ЦЕНТР ПО ПРОВЕДЕНИЮ СПАСАТЕЛЬНЫХ",                            {"bold": True, "size": 8,  "align": "center"}),
    ("ОПЕРАЦИЙ ОСОБОГО РИСКА «ЛИДЕР»",                               {"bold": True, "size": 8,  "align": "center"}),
    ("(ФГКУ «ЦСООР «Лидер»)",                                        {"bold": True, "size": 8,  "align": "center"}),
    ("Музыкальный проезд, дом 4, строение 1",                        {"size": 7,  "align": "center", "space_before": True}),
    ("поселок завода Мосрентген, внутригородское",                   {"size": 7,  "align": "center"}),
    ("муниципальное образование – муниципальный округ Коммунарка,", {"size": 7,  "align": "center"}),
    ("г. Москва, 108820,",                                           {"size": 7,  "align": "center"}),
    ("Телефон: 8-(495)-424-00-33, 8-(495)-339-76-88",                {"size": 7,  "align": "center"}),
    ("Факс: 8-(495)-424-00-11",                                      {"size": 7,  "align": "center"}),
    ("E-mail: info@lider.mchs.gov.ru ; intranet: lider@mchs.ru",     {"size": 7,  "align": "center"}),
    ("___________________ № _________________",                      {"size": 8,  "align": "center", "space_before": True}),
    ("На №                  от                  .",                  {"size": 8,  "align": "center"}),
]

_HEADER_RIGHT = [
    ("МЧС России",                                          {"bold": True, "size": 10}),
    ("Врид директора Департамента",                         {"size": 10, "space_before": True}),
    ("информационных технологий и связи",                   {"size": 10}),
    ("полковнику",                                          {"size": 10, "space_before": True}),
    ("Стёпину Р.Ю.",                                        {"size": 10}),
]

# 22 колонки — точные названия и ширины.
# Альбомная A4: лист 29.7 см, поля по 0.7 см → доступно 28.3 см.
# Сумма ширин ниже = 26.7 см: таблица занимает всю ширину листа
# (как в оригинале), оставаясь с запасом 1.5 см до правого края.
_COL_HEADERS = [
    # (header_label, width_cm)
    ("№ п/п",                                 0.7),
    ("Наименование средств связи,\n"
     "вычислительной и оргтехники",           3.8),
    ("Ед. учета",                             0.8),
    ("Потреб-ность",                          1.0),
    ("Состояло на 1 января {prev} г.",        1.2),
    ("Прибыло",                               0.9),
    ("Убыло",                                 0.9),
    ("Всего",                                 0.9),
    ("В наличии исправной",                   1.5),
    ("В наличии современной",                 1.6),
    ("В наличии со сроком службы свыше",      1.6),
    ("В «НЗ»",                                0.8),
    ("На «ТД»",                               0.8),
    ("Подменный фонд 2-3 кат",                1.2),
    ("Резерв МЧС России",                     1.2),
    ("В т.ч. кап. ремонт",                    1.2),
    ("На МБ",                                 0.8),
    ("В т.ч. списано",                        1.0),
    ("В запасах центров",                     1.1),
    ("Укомплектованность в %",                1.1),
    ("Недостает или излишествует «-» / «+»",  1.1),
    ("Примечание",                            1.5),
]

# Маппинг колонок → ключи item'а. None — производное / только в шапке.
# Используется при заполнении строк позиций.
_COL_FIELD = [
    None,        # 1: №
    None,        # 2: name (отдельно)
    None,        # 3: ед.
    "required",  # 4
    "start",     # 5
    "arrived",   # 6
    "removed",   # 7
    "_total",    # 8 (производная)
    "working",   # 9
    "modern",    # 10
    "overdue",   # 11
    None,        # 12 «НЗ»  — не редактируется в MVP
    None,        # 13 «ТД»
    None,        # 14 Подменный
    None,        # 15 Резерв
    None,        # 16 кап. ремонт
    None,        # 17 МБ
    None,        # 18 списано
    "plus",      # 19 в запасах центров — текущее `plus` маппится сюда
    "_percent",  # 20
    "_diff",     # 21
    "note",      # 22
]


def _set_cell_size(cell, font_size: int, *, bold: bool = False,
                   italic: bool = False, align: str | None = None):
    """Применяет шрифт/выравнивание ко всем runs в ячейке."""
    for p in cell.paragraphs:
        if align == "left":   p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if align == "right":  p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in p.runs:
            r.font.size = Pt(font_size)
            if bold:   r.bold = True
            if italic: r.italic = True


def _add_text(parent, text: str, *, bold=False, size=10, align=None):
    p = parent.add_paragraph()
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if align == "right":  p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold: run.bold = True
    return p


def _build_letterhead(doc: Document, year: int):
    """Шапка письма: слева реквизиты ФГКУ, справа адресат."""
    from docx.shared import Pt as _Pt

    def _put(cell, lines, *, first_indent=False):
        """
        Заливает в ячейку строки бланка. Первый абзац уже создан Word'ом —
        используем его, далее add_paragraph(). Все строки сжаты:
        space_before/after = 0, line_spacing = 1.0.
        """
        first = True
        for text, fmt in lines:
            if first:
                # У свежесозданной ячейки уже есть один пустой абзац — берём его
                p = cell.paragraphs[0]
                p.text = text
                first = False
            else:
                p = cell.add_paragraph(text)

            pf = p.paragraph_format
            pf.space_before = _Pt(fmt.get("space_before") and 6 or 0)
            pf.space_after  = _Pt(0)
            pf.line_spacing = 1.0

            if fmt.get("align") == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for r in p.runs:
                r.font.size = _Pt(fmt.get("size", 9))
                if fmt.get("bold"): r.bold = True

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells
    left.width  = Cm(11.5)
    right.width = Cm(15.0)

    _put(left,  _HEADER_LEFT)
    _put(right, _HEADER_RIGHT)

    # Убираем границы у этой служебной таблицы (она нужна только для верстки)
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _Ox
    tbl = table._tbl
    tblBorders = _Ox("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = _Ox(f"w:{edge}")
        b.set(_qn("w:val"), "nil")
        tblBorders.append(b)
    tbl.tblPr.append(tblBorders)

    # «Форма 3/СВЯЗЬ» в правом верхнем углу под адресатом —
    # короткий компактный абзац, сразу под letterhead.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = _Pt(0)
    p.paragraph_format.space_after  = _Pt(0)
    run = p.add_run("Форма 3/СВЯЗЬ")
    run.bold = True; run.font.size = _Pt(10)


def _build_main_header(table) -> None:
    """
    4-строчная шапка главной таблицы. Использует merge ячеек для
    группировки «Движение за отчётный период», «Состоит в наличии…»,
    «В том числе» и т.п. — повторяет визуал docx-шаблона.
    """
    # Создадим 4 строки шапки
    for _ in range(3):
        table.add_row()

    rows = table.rows  # 4 первые строки — шапка
    # Для краткости: r[i][j] = ячейка
    r = [list(row.cells) for row in rows[:4]]

    # Колонки 1-5 — vertical merge через все 4 строки
    for col_idx in (0, 1, 2, 3, 4):
        merged = r[0][col_idx]
        for k in range(1, 4):
            merged = merged.merge(r[k][col_idx])

    # Колонка-группа «Движение за отчётный период» (cols 6-7) — row 1 объединена,
    # затем row 2-4 у каждой свой подпис.
    movement = r[0][5].merge(r[0][6])  # row 1: «Движение …»

    # Колонка-группа «Состоит в наличии … кроме того» (cols 8-19) — row 1 merged
    in_stock = r[0][7]
    for c in range(8, 19):
        in_stock = in_stock.merge(r[0][c])

    # Cols 20-22 — vertical merge
    for col_idx in (19, 20, 21):
        merged = r[0][col_idx]
        for k in range(1, 4):
            merged = merged.merge(r[k][col_idx])

    # Row 2: «Прибыло» (col 6, merge через rows 2-4), «Убыло» (col 7), «Всего» (col 8 merge через 2-4),
    # «В том числе» (cols 9-19 объединены в row 2)
    pribilo = r[1][5]
    for k in range(2, 4):
        pribilo = pribilo.merge(r[k][5])
    ubilo = r[1][6]
    for k in range(2, 4):
        ubilo = ubilo.merge(r[k][6])
    total_h = r[1][7]
    for k in range(2, 4):
        total_h = total_h.merge(r[k][7])

    in_some = r[1][8]
    for c in range(9, 19):
        in_some = in_some.merge(r[1][c])

    # Row 3: каждая клетка cols 9-19 — отдельный подпис, vertical merge через rows 3-4
    for c in range(8, 19):
        r[2][c].merge(r[3][c])

    # Заполняем тексты
    # Row 1 (top)
    r[0][0].text = "№ п/п"
    r[0][1].text = "Наименование средств связи, вычислительной и оргтехники"
    r[0][2].text = "Ед. учета"
    r[0][3].text = "Потреб-ность"
    r[0][4].text = "Состояло на 1 января {} г.".format(0)  # заполним после

    movement.text = "Движение за отчётный период"
    in_stock.text = "Состоит в наличии на 1 января {} г.\nкроме того".format(0)

    r[0][19].text = "Укомплек-тованность в %"
    r[0][20].text = "Недостает или излишествует «-» или «+»"
    r[0][21].text = "Примечание"

    # Row 2: уточняющие подписи
    pribilo.text = "Прибыло"
    ubilo.text   = "Убыло"
    total_h.text = "Всего"
    in_some.text = "В том числе"

    # Row 3: заголовки cols 9-19
    r[2][8].text  = "В наличии исправной"
    r[2][9].text  = "В наличии современной согласно Методике расчёта"
    r[2][10].text = "В наличии со сроком службы свыше установленного"
    r[2][11].text = "В «НЗ»"
    r[2][12].text = "На «ТД»"
    r[2][13].text = "Подменный фонд 2-3 кат"
    r[2][14].text = "Резерв МЧС России"
    r[2][15].text = "В т.ч. кап. ремонт"
    r[2][16].text = "На МБ"
    r[2][17].text = "В т.ч. списано"
    r[2][18].text = "В запасах центров"

    # Row 4: нумерация колонок 1..22
    for i in range(22):
        r[3][i].text = str(i + 1)

    # Стилизуем все ячейки шапки
    for row in rows[:4]:
        for c in row.cells:
            _set_cell_size(c, 7, bold=True, align="center")


def _set_table_fixed_layout(table, widths_cm: list[float]):
    """
    Принудительно фиксирует ширины колонок: tblLayout=fixed + tblGrid с явными
    значениями. Без этого Word при открытии часто перерасчитывает колонки
    под видимый контент, и таблица «уезжает» за правый край листа.
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _Ox

    tbl = table._tbl
    tblPr = tbl.tblPr

    # 1. Зафиксировать layout
    layout = _Ox("w:tblLayout")
    layout.set(_qn("w:type"), "fixed")
    # Удалить существующий, если был
    for el in tblPr.findall(_qn("w:tblLayout")):
        tblPr.remove(el)
    tblPr.append(layout)

    # 2. Прописать tblGrid с шириной каждой колонки в DXA (1cm = 567 dxa)
    for el in tbl.findall(_qn("w:tblGrid")):
        tbl.remove(el)
    tblGrid = _Ox("w:tblGrid")
    for w in widths_cm:
        gc = _Ox("w:gridCol")
        gc.set(_qn("w:w"), str(int(round(w * 567))))
        tblGrid.append(gc)
    # Вставить tblGrid сразу после tblPr
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)

    # 3. Дополнительно — явные width на каждую ячейку
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _build_main_table(doc: Document, year: int, data: list[dict]):
    """Главная таблица 22 колонки + строки направлений и позиций."""
    table = doc.add_table(rows=1, cols=22)
    table.style = "Table Grid"
    table.autofit = False

    widths = [w for _, w in _COL_HEADERS]
    _set_table_fixed_layout(table, widths)

    _build_main_header(table)

    # Заполняем динамические части шапки (год)
    rows = table.rows
    # row 0 col 4: Состояло на 1 января {year-1} г.
    rows[0].cells[4].text = f"Состояло на 1 января {year - 1} г."
    # row 0 in_stock: Состоит в наличии на 1 января {year} г. кроме того
    # in_stock — это ячейка row[0][7] после merge, текст уже задан, переписать
    rows[0].cells[7].text = f"Состоит в наличии на 1 января {year} г. кроме того"
    for c in (rows[0].cells[4], rows[0].cells[7]):
        _set_cell_size(c, 7, bold=True, align="center")

    # Маппинг полей item → индекс колонки docx (0-based, итого 22 колонки):
    #   3 required, 4 start, 5 arrived, 6 removed, 7 total
    #   8 working, 9 modern, 10 overdue
    #   11 nz, 12 td, 13 backup_fund, 14 mchs_reserve,
    #   15 capital_repair, 16 mb, 17 written_off, 18 plus (=в запасах центров)
    #   19 %, 20 +/-, 21 note
    _SUM_FIELDS = (
        "required", "start", "arrived", "removed",
        "working", "modern", "overdue",
        "nz", "td", "backup_fund", "mchs_reserve",
        "capital_repair", "mb", "written_off", "plus",
    )

    # Строки данных — итог направления + позиции
    for cat in data:
        # ── Сводная строка направления ──────────────────────────────────
        leaf_items = [it for it in cat.get("items", []) if not it.get("is_group")]
        cat_sum = {f: 0 for f in _SUM_FIELDS}
        for it in leaf_items:
            d = compute_derived(it)
            for f in cat_sum:
                cat_sum[f] += d.get(f, 0)
        total_sum   = cat_sum["start"] + cat_sum["arrived"] - cat_sum["removed"]
        req_sum     = cat_sum["required"]
        percent_sum = round(total_sum / req_sum * 100) if req_sum > 0 \
                      else (0 if total_sum == 0 else 100)
        diff_sum    = total_sum - req_sum

        def _s(n: int) -> str:
            """Выводим число только если оно ненулевое — пустые ячейки выглядят аккуратнее."""
            return str(n) if n else ""

        cells = table.add_row().cells
        cells[0].text  = str(cat.get("index", ""))
        cells[1].text  = cat.get("title", "")
        cells[2].text  = cat.get("unit", "")
        cells[3].text  = _s(cat_sum["required"])
        cells[4].text  = _s(cat_sum["start"])
        cells[5].text  = _s(cat_sum["arrived"])
        cells[6].text  = _s(cat_sum["removed"])
        cells[7].text  = _s(total_sum)
        cells[8].text  = _s(cat_sum["working"])
        cells[9].text  = _s(cat_sum["modern"])
        cells[10].text = _s(cat_sum["overdue"])
        cells[11].text = _s(cat_sum["nz"])
        cells[12].text = _s(cat_sum["td"])
        cells[13].text = _s(cat_sum["backup_fund"])
        cells[14].text = _s(cat_sum["mchs_reserve"])
        cells[15].text = _s(cat_sum["capital_repair"])
        cells[16].text = _s(cat_sum["mb"])
        cells[17].text = _s(cat_sum["written_off"])
        cells[18].text = _s(cat_sum["plus"])
        cells[19].text = str(percent_sum)
        cells[20].text = f"{'+' if diff_sum > 0 else ''}{diff_sum}" if diff_sum != 0 else "0"

        for c in cells:
            _set_cell_size(c, 7, bold=True, align="center")
        _set_cell_size(cells[1], 7, bold=True, align="left")

        # ── Строки позиций ──────────────────────────────────────────────
        for it in cat.get("items", []):
            d = compute_derived(it)
            cells = table.add_row().cells
            cells[0].text = ""
            cells[1].text = (" " * 2 if it.get("is_group") else " " * 4) + d["name"]
            cells[2].text = ""

            if it.get("is_group"):
                for c in cells[3:]:
                    c.text = ""
                _set_cell_size(cells[1], 7, bold=True, italic=True, align="left")
                continue

            cells[3].text  = _s(d["required"])
            cells[4].text  = _s(d["start"])
            cells[5].text  = _s(d["arrived"])
            cells[6].text  = _s(d["removed"])
            cells[7].text  = _s(d["total"])
            cells[8].text  = _s(d["working"])
            cells[9].text  = _s(d["modern"])
            cells[10].text = _s(d["overdue"])
            cells[11].text = _s(d.get("nz", 0))
            cells[12].text = _s(d.get("td", 0))
            cells[13].text = _s(d.get("backup_fund", 0))
            cells[14].text = _s(d.get("mchs_reserve", 0))
            cells[15].text = _s(d.get("capital_repair", 0))
            cells[16].text = _s(d.get("mb", 0))
            cells[17].text = _s(d.get("written_off", 0))
            cells[18].text = _s(d["plus"])
            cells[19].text = str(d["percent"])
            cells[20].text = f"{'+' if d['diff'] > 0 else ''}{d['diff']}" if d['diff'] != 0 else "0"
            cells[21].text = d.get("note", "") or ""

            for c in cells:
                _set_cell_size(c, 7, align="center")
            _set_cell_size(cells[1], 7, align="left")


def _build_signature(doc: Document):
    """Подпись начальника + контактное лицо."""
    doc.add_paragraph()  # отступ
    p = doc.add_paragraph()
    run = p.add_run("Начальник ФГКУ «ЦСООР «Лидер»")
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("генерал-майор")
    run.font.size = Pt(11)
    run2 = p.add_run("\t\t\t\t\t\t\t\tА.А. Саввин")
    run2.font.size = Pt(11)

    p = doc.add_paragraph("__.01.{}".format(2026))
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph("Саноян Арман Норикович")
    p.runs[0].font.size = Pt(9)
    p = doc.add_paragraph("8(495)339-76-88 (169)")
    p.runs[0].font.size = Pt(9)


def _build_appendix(doc: Document, year: int, data: list[dict]):
    """Приложение к форме 3/СВЯЗЬ — обобщённая таблица 10 колонок."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Приложение к форме 3/СВЯЗЬ")
    run.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ОБОБЩЁННЫЕ СВЕДЕНИЯ")
    run.bold = True; run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("о наличии и обеспеченности средствами связи, вычислительной и оргтехникой")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"(по состоянию на 1.01.{year} г.)")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Таблица: 10 колонок (№ п/п, Наименование, Ед.учета, По штату, Имеется ВСЕГО,
    # В наличии исправной, В наличии современной, В наличии со сроком, +/-, %, Примечание)
    APX_HEADERS = [
        "№ п/п", "Наименование средств связи, вычислительной и оргтехники",
        "Ед. учёта", "Положено по штату", f"Имеется в наличии на 1.01.{year} г. ВСЕГО",
        "В наличии исправной", "В наличии современной",
        "В наличии со сроком службы свыше установленного приказом МЧС России от 25.11.2016 № 624",
        "Недостает («-») или излишествует («+»)", "Укомплектованность в %",
        "Примечание",
    ]

    table = doc.add_table(rows=2, cols=len(APX_HEADERS))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(APX_HEADERS):
        hdr[i].text = h
        _set_cell_size(hdr[i], 8, bold=True, align="center")
    # Row 2: нумерация
    for i, c in enumerate(table.rows[1].cells):
        c.text = str(i + 1)
        _set_cell_size(c, 8, align="center")

    # Данные: одна строка на направление (агрегаты)
    for idx, cat in enumerate(data, start=1):
        leafs = [it for it in cat.get("items", []) if not it.get("is_group")]
        s = {f: 0 for f in ("required", "start", "arrived", "removed",
                            "working", "modern", "overdue")}
        for it in leafs:
            d = compute_derived(it)
            for f in s:
                s[f] += d[f]
        total = s["start"] + s["arrived"] - s["removed"]
        diff  = total - s["required"]
        pct   = round(total / s["required"] * 100) if s["required"] > 0 \
                else (0 if total == 0 else 100)

        cells = table.add_row().cells
        cells[0].text = f"{idx}."
        cells[1].text = cat.get("title", "")
        cells[2].text = cat.get("unit", "")
        cells[3].text = str(s["required"]) if s["required"] else ""
        cells[4].text = str(total)         if total         else ""
        cells[5].text = str(s["working"])  if s["working"]  else ""
        cells[6].text = str(s["modern"])   if s["modern"]   else ""
        cells[7].text = str(s["overdue"])  if s["overdue"]  else ""
        cells[8].text = f"{'+' if diff > 0 else ''}{diff}" if diff != 0 else "0"
        cells[9].text = str(pct)
        cells[10].text = ""

        for c in cells:
            _set_cell_size(c, 8, align="center")
        _set_cell_size(cells[1], 8, align="left")

    # Подпись после приложения — та же
    _build_signature(doc)


def _build_docx(unit_username: str, year: int, data: list[dict]) -> io.BytesIO:
    """
    Строит .docx Форма 3-СВЯЗЬ один-в-один с шаблоном:
      • бланк ФГКУ «ЦСООР «Лидер» + адресат МЧС России (Стёпину Р.Ю.)
      • заголовок «Форма 3/СВЯЗЬ» + «ДОНЕСЕНИЕ…»
      • главная таблица 22 колонки с многоуровневой шапкой
      • подпись (А.А. Саввин) и контакт (Саноян А.Н.)
      • приложение: «ОБОБЩЁННЫЕ СВЕДЕНИЯ» — 10-колоночная сводная таблица
    """
    doc = Document()

    # A4 альбомная — таблица в 22 колонки иначе не помещается.
    # Поля сужены до 0.7 см: содержательная область ~28.3 см, чтобы влезла
    # 22-колоночная таблица суммарной ~24 см + запас на padding.
    section = doc.sections[0]
    from docx.enum.section import WD_ORIENT
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin   = Cm(0.7)
    section.right_margin  = Cm(0.7)

    # 1. Бланк (адресатура слева, получатель справа)
    _build_letterhead(doc, year)

    # 2. Заголовок «ДОНЕСЕНИЕ»
    doc.add_paragraph()
    _add_text(doc, "ДОНЕСЕНИЕ", bold=True, size=13, align="center")
    _add_text(doc,
        "о потребности, наличии, движении и качественном состоянии "
        "средств связи, вычислительной и оргтехники",
        size=11, align="center",
    )
    _add_text(doc, "ФГКУ «ЦСООР «Лидер»",            size=11, align="center")
    _add_text(doc, "(наименование учреждения, органа управления МЧС России)",
              size=8,  align="center")
    _add_text(doc, f"По состоянию на 1 января {year} г.", size=10, align="right")

    # 3. Главная 22-колоночная таблица
    _build_main_table(doc, year, data)

    # 4. Подпись
    _build_signature(doc)

    # 5. Приложение (новая страница)
    _build_appendix(doc, year, data)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.get("/export")
def export_report(
        year:         int              = Query(..., ge=2020, le=2100),
        unit:         str | None       = Query(None, max_length=100),
        db:           Session          = Depends(get_db),
        current_user: User             = Depends(get_current_user),
):
    unit_name = _resolve_unit(current_user, unit)
    rep = _get_or_create_report(db, unit_name, year)

    buf = _build_docx(unit_name, year, rep.data)

    filename = f"Форма_3-СВЯЗЬ_{year}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition":
                f'attachment; filename*=UTF-8\'\'{quote(filename)}',
        },
    )
