# standalone_maps/server.py
"""
Автономный сервер «Карты зон ответственности» — БЕЗ логинов и учёток.

Разворачивает на одном компьютере (с интернетом) две карты:
  • «Карта ОД»   — базовая точка, поиск адреса, маршрут, зоны ответственности
  • «Карта зон»  — импорт координат из Excel/Word, пересчёт МСК-77 → WGS-84,
                   таблица координат, подвижка контура, экспорт JPG/PDF (1:2000)

Всё «само разворачивается»:
  • БД — SQLite (файл _data/maps.db рядом, создаётся при первом старте)
  • тайлы карты — OpenStreetMap (без API-ключа), проксируются локально
    (same-origin, чтобы работал экспорт картинки)
  • геокодер — Nominatim (OSM), маршруты — OSRM, оба без ключей
  • авторизации НЕТ: открыл http://127.0.0.1:8077 и работаешь

Запуск:
    python server.py                # стартует сервер и открывает браузер
    python server.py --no-browser   # без авто-открытия
    PORT=9000 python server.py      # другой порт
"""

import io
import json
import math
import os
import re
import sqlite3
import sys
import threading
import webbrowser
from contextlib import closing
from pathlib import Path
from typing import List, Optional

import httpx
from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, FastAPI, File, Form, HTTPException, Response, UploadFile
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook, load_workbook
from pydantic import BaseModel, Field

import geodesy


# ─── Пути (работают и как скрипт, и как PyInstaller .exe) ─────────────────────

def _app_dir() -> Path:
    if getattr(sys, "frozen", False):          # собранный .exe
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent

def _static_dir() -> Path:
    # PyInstaller onefile распаковывает данные в sys._MEIPASS
    base = getattr(sys, "_MEIPASS", None)
    if base:
        p = Path(base) / "static"
        if p.exists():
            return p
    return _app_dir() / "static"

APP_DIR    = _app_dir()
DATA_DIR   = APP_DIR / "_data"
DATA_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH    = DATA_DIR / "maps.db"
TILE_CACHE = DATA_DIR / "tiles"
STATIC_DIR = _static_dir()

OSM_TILE_URL = "https://tile.openstreetmap.org/{z}/{x}/{y}.png"
NOMINATIM    = "https://nominatim.openstreetmap.org/search"
OSRM_BASE    = "https://router.project-osrm.org"
USER_AGENT   = "PODS2-StandaloneMaps/1.0 (local single-PC deployment)"

YANDEX_TILE    = "https://core-renderer-tiles.maps.yandex.net/tiles"
YANDEX_GEOCODE = "https://geocode-maps.yandex.ru/1.x/"
YANDEX_SUGGEST = "https://suggest-maps.yandex.ru/v1/suggest"
# Bbox Москва+МО для приоритета геокодера/подсказок (как в исходной «Карте ОД»).
MO_BBOX  = "35.15,54.25~40.20,56.97"
MO_LL    = "37.6173,55.7558"
MO_SPN   = "5,3"


def _load_yandex_key() -> str:
    """Ключ Яндекс.Карт: из env YANDEX_MAPS_API_KEY или файла рядом (yandex.key)."""
    k = (os.environ.get("YANDEX_MAPS_API_KEY") or "").strip()
    if k:
        return k
    for name in ("yandex.key", "yandex_key.txt"):
        p = APP_DIR / name
        if p.exists():
            try:
                t = p.read_text(encoding="utf-8").strip()
                if t:
                    return t
            except OSError:
                pass
    return ""


YANDEX_KEY = _load_yandex_key()
PROVIDER   = "yandex" if YANDEX_KEY else "osm"   # есть ключ → Яндекс, иначе OSM

MAPS = ("od", "zone")   # допустимые карты


# ─── База данных (SQLite, без внешних зависимостей) ───────────────────────────

def _db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def _init_db() -> None:
    with closing(_db()) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS zones (
                id              INTEGER PRIMARY KEY AUTOINCREMENT,
                map             TEXT NOT NULL,
                name            TEXT NOT NULL,
                role            TEXT,
                color           TEXT NOT NULL DEFAULT '#1976d2',
                points_json     TEXT NOT NULL DEFAULT '[]',
                src_points_json TEXT,
                coord_system    TEXT,
                sort_order      INTEGER NOT NULL DEFAULT 0
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS od_settings (
                id           INTEGER PRIMARY KEY CHECK (id = 1),
                base_address TEXT,
                base_lat     REAL,
                base_lng     REAL
            )
        """)
        conn.execute("INSERT OR IGNORE INTO od_settings (id) VALUES (1)")
        conn.commit()


def _zone_row_to_dict(r: sqlite3.Row) -> dict:
    return {
        "id":           r["id"],
        "name":         r["name"],
        "role":         r["role"],
        "color":        r["color"],
        "points":       json.loads(r["points_json"] or "[]"),
        "src_points":   json.loads(r["src_points_json"]) if r["src_points_json"] else None,
        "coord_system": r["coord_system"],
        "sort_order":   r["sort_order"],
    }


# ─── Pydantic ─────────────────────────────────────────────────────────────────

class ZoneIn(BaseModel):
    name:       str = Field(..., min_length=1, max_length=200)
    role:       Optional[str] = Field(default=None, max_length=200)
    color:      str = Field(default="#1976d2", max_length=20)
    points:     list = Field(default_factory=list)
    sort_order: int = 0

class ZonePatch(BaseModel):
    name:       Optional[str] = Field(default=None, min_length=1, max_length=200)
    role:       Optional[str] = Field(default=None, max_length=200)
    color:      Optional[str] = Field(default=None, max_length=20)
    points:     Optional[list] = None
    sort_order: Optional[int] = None

class SettingsIn(BaseModel):
    base_address: Optional[str] = Field(default=None, max_length=500)
    base_lat:     Optional[float] = None
    base_lng:     Optional[float] = None

class RouteIn(BaseModel):
    src: List[float] = Field(..., min_length=2, max_length=2)
    dst: List[float] = Field(..., min_length=2, max_length=2)


def _clean_points(raw: list) -> list:
    out = []
    for p in raw or []:
        try:
            lat, lng = float(p[0]), float(p[1])
        except (TypeError, ValueError, IndexError):
            continue
        if -90.0 <= lat <= 90.0 and -180.0 <= lng <= 180.0:
            out.append([round(lat, 7), round(lng, 7)])
    return out

def _check_map(m: str) -> str:
    if m not in MAPS:
        raise HTTPException(status_code=404, detail="Неизвестная карта")
    return m


app = FastAPI(title="Карты зон ответственности (автономно)")
api = APIRouter(prefix="/api")


# ─── Зоны: CRUD (общий для обеих карт) ───────────────────────────────────────

@api.get("/{map}/zones")
def list_zones(map: str):
    _check_map(map)
    with closing(_db()) as conn:
        rows = conn.execute(
            "SELECT * FROM zones WHERE map = ? ORDER BY sort_order, id", (map,)
        ).fetchall()
    return [_zone_row_to_dict(r) for r in rows]

@api.post("/{map}/zones", status_code=201)
def create_zone(map: str, payload: ZoneIn):
    _check_map(map)
    with closing(_db()) as conn:
        cur = conn.execute(
            "INSERT INTO zones (map,name,role,color,points_json,sort_order) "
            "VALUES (?,?,?,?,?,?)",
            (map, payload.name.strip(), (payload.role or "").strip() or None,
             payload.color or "#1976d2", json.dumps(_clean_points(payload.points)),
             payload.sort_order or 0),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM zones WHERE id = ?", (cur.lastrowid,)).fetchone()
    return _zone_row_to_dict(row)

@api.patch("/{map}/zones/{zone_id}")
def patch_zone(map: str, zone_id: int, payload: ZonePatch):
    _check_map(map)
    with closing(_db()) as conn:
        row = conn.execute("SELECT * FROM zones WHERE id = ? AND map = ?", (zone_id, map)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Зона не найдена")
        name  = payload.name.strip() if payload.name is not None else row["name"]
        role  = ((payload.role or "").strip() or None) if payload.role is not None else row["role"]
        color = payload.color if payload.color is not None else row["color"]
        pts   = json.dumps(_clean_points(payload.points)) if payload.points is not None else row["points_json"]
        so    = payload.sort_order if payload.sort_order is not None else row["sort_order"]
        conn.execute(
            "UPDATE zones SET name=?, role=?, color=?, points_json=?, sort_order=? WHERE id=?",
            (name, role, color, pts, so, zone_id),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM zones WHERE id = ?", (zone_id,)).fetchone()
    return _zone_row_to_dict(row)

@api.delete("/{map}/zones/{zone_id}", status_code=204)
def delete_zone(map: str, zone_id: int):
    _check_map(map)
    with closing(_db()) as conn:
        conn.execute("DELETE FROM zones WHERE id = ? AND map = ?", (zone_id, map))
        conn.commit()

@api.delete("/{map}/zones", status_code=204)
def clear_zones(map: str):
    _check_map(map)
    with closing(_db()) as conn:
        conn.execute("DELETE FROM zones WHERE map = ?", (map,))
        conn.commit()


# ─── Базовая точка (только карта ОД) ─────────────────────────────────────────

@api.get("/od/settings")
def get_settings():
    with closing(_db()) as conn:
        r = conn.execute("SELECT * FROM od_settings WHERE id = 1").fetchone()
    return {"base_address": r["base_address"], "base_lat": r["base_lat"], "base_lng": r["base_lng"]}

@api.put("/od/settings")
def put_settings(payload: SettingsIn):
    with closing(_db()) as conn:
        conn.execute(
            "UPDATE od_settings SET base_address=?, base_lat=?, base_lng=? WHERE id=1",
            ((payload.base_address or "").strip() or None, payload.base_lat, payload.base_lng),
        )
        conn.commit()
    return get_settings()


# ─── Конфиг провайдера карт ───────────────────────────────────────────────────

@api.get("/config")
def config():
    # provider: yandex (если задан ключ) или osm. suggest — только у Яндекса.
    return {"provider": PROVIDER, "suggest": PROVIDER == "yandex",
            "attribution": "© Яндекс" if PROVIDER == "yandex" else "© OpenStreetMap"}


# ─── Геокодер: Яндекс (если есть ключ) или Nominatim ─────────────────────────

def _ll_from_uri(uri: str):
    if not uri:
        return None
    import urllib.parse as _u
    m = re.search(r"[?&]ll=([0-9.\-]+),([0-9.\-]+)", _u.unquote(uri))
    if not m:
        return None
    try:
        return float(m.group(2)), float(m.group(1))   # uri даёт lng,lat → возвращаем lat,lng
    except ValueError:
        return None


@api.get("/geocode")
async def geocode(q: str = "", uri: str = ""):
    q = (q or "").strip()
    uri = (uri or "").strip()
    if not q and not uri:
        raise HTTPException(status_code=400, detail="Пустой запрос")

    if PROVIDER == "yandex":
        params = {"apikey": YANDEX_KEY, "format": "json", "lang": "ru_RU", "results": "15"}
        if uri:
            # Точный объект из подсказки (организация/адрес) — геокодер вернёт ровно его.
            params["uri"] = uri
        else:
            ql = q.lower()
            geo = q if any(w in ql for w in ("москва","московская","подмосков","мо ","мо,")) \
                     else f"Россия, Москва, {q}"
            params["geocode"] = geo
            params["bbox"] = MO_BBOX
        try:
            async with httpx.AsyncClient(timeout=12.0) as c:
                r = await c.get(YANDEX_GEOCODE, params=params)
                r.raise_for_status()
                data = r.json()
        except httpx.HTTPError:
            raise HTTPException(status_code=502, detail="Геокодер Яндекса недоступен")
        out = []
        try:
            members = data["response"]["GeoObjectCollection"]["featureMember"]
        except (KeyError, TypeError):
            members = []
        for m in members:
            try:
                obj = m["GeoObject"]
                lng, lat = (float(v) for v in obj["Point"]["pos"].split())
                meta = obj.get("metaDataProperty", {}).get("GeocoderMetaData", {})
                out.append({"text": meta.get("text") or obj.get("name") or "",
                            "lat": lat, "lng": lng, "kind": meta.get("kind", "")})
            except (KeyError, ValueError, TypeError, IndexError):
                continue
        return {"results": out}

    # OSM fallback (Nominatim)
    params = {"q": q, "format": "jsonv2", "limit": "10", "accept-language": "ru"}
    try:
        async with httpx.AsyncClient(timeout=12.0, headers={"User-Agent": USER_AGENT}) as c:
            r = await c.get(NOMINATIM, params=params)
            r.raise_for_status()
            data = r.json()
    except httpx.HTTPError:
        raise HTTPException(status_code=502, detail="Геокодер недоступен")
    out = []
    for it in data:
        try:
            out.append({"text": it.get("display_name", ""), "lat": float(it["lat"]),
                        "lng": float(it["lon"]), "kind": it.get("type", "")})
        except (KeyError, ValueError, TypeError):
            continue
    return {"results": out}


@api.get("/suggest")
async def suggest(q: str):
    """Подсказки адресов при наборе (только Яндекс). Без ключа — пусто."""
    q = (q or "").strip()
    if PROVIDER != "yandex" or len(q) < 2:
        return {"results": []}
    params = {"apikey": YANDEX_KEY, "text": q, "lang": "ru_RU", "results": "10",
              "ll": MO_LL, "spn": MO_SPN, "print_address": "1", "attrs": "uri", "types": "geo,biz"}
    try:
        async with httpx.AsyncClient(timeout=8.0) as c:
            r = await c.get(YANDEX_SUGGEST, params=params)
            if r.status_code in (401, 403):
                return {"results": []}     # ключ без права на Suggest — тихо деградируем
            r.raise_for_status()
            data = r.json()
    except httpx.HTTPError:
        return {"results": []}
    out = []
    for it in data.get("results", []):
        try:
            title = (it.get("title") or {}).get("text") or ""
            if not title:
                continue
            sub = (it.get("subtitle") or {}).get("text") or ""
            addr = ((it.get("address") or {}).get("formatted_address") or "")
            uri = it.get("uri") or ""
            ll = _ll_from_uri(uri)
            out.append({"title": title, "subtitle": sub or addr, "uri": uri,
                        "lat": ll[0] if ll else None, "lng": ll[1] if ll else None})
        except (KeyError, TypeError):
            continue
    return {"results": out}

@api.post("/route")
async def route(payload: RouteIn):
    s_lat, s_lng = payload.src
    d_lat, d_lng = payload.dst
    url = (f"{OSRM_BASE}/route/v1/driving/{s_lng},{s_lat};{d_lng},{d_lat}"
           "?overview=full&geometries=geojson&steps=false")
    try:
        async with httpx.AsyncClient(timeout=15.0, headers={"User-Agent": USER_AGENT}) as c:
            r = await c.get(url)
            r.raise_for_status()
            data = r.json()
    except httpx.HTTPError:
        raise HTTPException(status_code=502, detail="Маршрутизатор недоступен")
    if data.get("code") != "Ok" or not data.get("routes"):
        raise HTTPException(status_code=400, detail="Маршрут не построен")
    rt = data["routes"][0]
    return {"geometry": rt.get("geometry"), "distance_m": rt.get("distance"), "duration_s": rt.get("duration")}


# ─── Импорт координат (карта зон): Excel / Word + МСК-77 → WGS-84 ─────────────

_PALETTE = ["#1976d2","#e53935","#43a047","#fb8c00","#8e24aa","#00897b","#c2185b",
            "#3949ab","#7cb342","#f4511e","#00acc1","#6d4c41","#d81b60","#5e35b1","#039be5"]
_HEX_RE = re.compile(r"^#(?:[0-9a-fA-F]{3}|[0-9a-fA-F]{6})$")
_LAT_KEYS  = ("широта","шир","lat","latitude")
_LNG_KEYS  = ("долгота","долг","lng","lon","long","longitude")
_X_KEYS    = ("x","север","north","абсцисса")
_Y_KEYS    = ("y","восток","east","ордината")
_ZONE_KEYS = ("зона","зоны","участок","группа","group","zone","наименование","название","объект","район","name")
_ROLE_KEYS = ("роль","категория","тип","role","подпись","примечание","комментарий")
_COLOR_KEYS = ("цвет","color","colour")
_COMBO_KEYS = ("координаты","коорд","coords","coordinates","latlng","ll")

def _norm_header(v) -> str:
    return re.sub(r"[^a-zа-я0-9]", "", str(v or "").strip().lower())

def _match_col(header, keys) -> bool:
    h = _norm_header(header)
    return bool(h) and any(_norm_header(k) in h for k in keys)

def _to_float(v):
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    m = re.search(r"-?\d+(?:\.\d+)?", str(v).strip().replace(",", "."))
    return float(m.group(0)) if m else None

def _split_combo(v):
    if v is None:
        return None
    nums = [n for n in (_to_float(p) for p in re.split(r"[;,\s]+", str(v).strip()) if p) if n is not None]
    return (nums[0], nums[1]) if len(nums) >= 2 else None

def _orient(a, b):
    if abs(a) > 90 >= abs(b):
        return b, a
    return a, b

def _cell(row, idx):
    return row[idx] if idx is not None and idx < len(row) else None

def _extract_rows_xlsx(data: bytes):
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=400, detail=f"Не удалось прочитать Excel: {exc}")
    ws = wb.active
    rows = [list(r) for r in ws.iter_rows(values_only=True)]
    return [r for r in rows if any(c is not None and str(c).strip() != "" for c in r)]

def _extract_rows_docx(data: bytes):
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=400, detail=f"Не удалось прочитать Word: {exc}")
    if not doc.tables:
        raise HTTPException(status_code=400, detail="В документе нет таблицы с координатами")
    rows = []
    for row in doc.tables[0].rows:
        cells = [c.text.strip() for c in row.cells]
        if any(cells):
            rows.append(cells)
    return rows

def _detect_columns(rows, coord_system):
    is_msk = coord_system.startswith("msk")
    found = {"lat":None,"lng":None,"x":None,"y":None,"zone":None,"role":None,"color":None,"combo":None}
    header_row = -1
    for i, row in enumerate(rows[:6]):
        cand = dict.fromkeys(found, None)
        for ci, cell in enumerate(row):
            if is_msk:
                if cand["x"] is None and _match_col(cell, _X_KEYS): cand["x"] = ci
                elif cand["y"] is None and _match_col(cell, _Y_KEYS): cand["y"] = ci
            else:
                if cand["lat"] is None and _match_col(cell, _LAT_KEYS): cand["lat"] = ci
                elif cand["lng"] is None and _match_col(cell, _LNG_KEYS): cand["lng"] = ci
                elif cand["combo"] is None and _match_col(cell, _COMBO_KEYS): cand["combo"] = ci
            if cand["zone"] is None and _match_col(cell, _ZONE_KEYS): cand["zone"] = ci
            if cand["role"] is None and _match_col(cell, _ROLE_KEYS): cand["role"] = ci
            if cand["color"] is None and _match_col(cell, _COLOR_KEYS): cand["color"] = ci
        ok = (cand["x"] is not None and cand["y"] is not None) if is_msk else \
             ((cand["lat"] is not None and cand["lng"] is not None) or cand["combo"] is not None)
        if ok:
            found = cand; header_row = i; break
    if header_row >= 0:
        data_rows = rows[header_row + 1:]
    else:
        ncols = max((len(r) for r in rows), default=0)
        first_nums = [c for c in (rows[0] if rows else []) if _to_float(c) is not None]
        data_rows = rows if len(first_nums) >= 2 else rows[1:]
        if ncols >= 3:
            found["zone"] = 0
            if is_msk: found["x"], found["y"] = 1, 2
            else: found["lat"], found["lng"] = 1, 2
        else:
            if is_msk: found["x"], found["y"] = 0, 1
            else: found["lat"], found["lng"] = 0, 1
    return found, data_rows

def _rows_to_zones(rows, coord_system, default_name):
    rows = [r for r in rows if any(c is not None and str(c).strip() != "" for c in r)]
    if not rows:
        raise HTTPException(status_code=400, detail="Файл пустой")
    is_msk = coord_system.startswith("msk")
    cols, data_rows = _detect_columns(rows, coord_system)
    groups, order, skipped = {}, [], 0
    for row in data_rows:
        lat = lng = src = None
        if is_msk:
            x = _to_float(_cell(row, cols["x"]))
            y = _to_float(_cell(row, cols["y"]))
            if x is None or y is None or abs(x) < 100:
                skipped += 1; continue
            try:
                lat, lng = geodesy.msk77_to_wgs84(x, y, coord_system)
            except Exception:  # noqa: BLE001
                skipped += 1; continue
            src = [x, y]
        else:
            if cols["combo"] is not None:
                pair = _split_combo(_cell(row, cols["combo"]))
                if pair: lat, lng = _orient(pair[0], pair[1])
            if lat is None or lng is None:
                a = _to_float(_cell(row, cols["lat"]))
                b = _to_float(_cell(row, cols["lng"]))
                if a is None or b is None:
                    skipped += 1; continue
                lat, lng = _orient(a, b)
        if not (-90.0 <= lat <= 90.0 and -180.0 <= lng <= 180.0):
            skipped += 1; continue
        zname = _cell(row, cols["zone"])
        zname = str(zname).strip() if zname is not None and str(zname).strip() else default_name
        if zname not in groups:
            groups[zname] = {"points": [], "src": [], "color": None, "role": None}
            order.append(zname)
        groups[zname]["points"].append([round(lat,7), round(lng,7)])
        groups[zname]["src"].append(src if src is not None else [round(lat,7), round(lng,7)])
        cval = _cell(row, cols["color"])
        if cval and groups[zname]["color"] is None:
            cs = str(cval).strip(); cs = cs if cs.startswith("#") else f"#{cs}"
            if _HEX_RE.match(cs): groups[zname]["color"] = cs
        rval = _cell(row, cols["role"])
        if rval and groups[zname]["role"] is None:
            groups[zname]["role"] = str(rval).strip()
    if not order:
        hint = "колонки X и Y (МСК-77)" if is_msk else "колонки с широтой и долготой"
        raise HTTPException(status_code=400,
            detail=f"Не нашли ни одной валидной координаты. Проверьте, что есть {hint} (или скачайте шаблон).")
    parsed = []
    for i, name in enumerate(order):
        g = groups[name]
        parsed.append({"name": name, "role": g["role"],
                       "color": g["color"] or _PALETTE[i % len(_PALETTE)],
                       "points": g["points"], "src_points": g["src"]})
    return parsed, skipped

def _filename_stem(filename: str) -> str:
    base = (filename or "").rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    return (base.rsplit(".", 1)[0].strip()[:200]) or "Импорт"

@api.post("/zone/import")
async def import_zones(file: UploadFile = File(...), mode: str = Form("replace"),
                       coord_system: str = Form("wgs84")):
    name = (file.filename or "").lower()
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Файл пустой")
    if name.endswith((".xlsx", ".xlsm")):
        rows = _extract_rows_xlsx(data)
    elif name.endswith(".docx"):
        rows = _extract_rows_docx(data)
    else:
        raise HTTPException(status_code=400, detail="Поддерживаются файлы .xlsx и .docx")
    parsed, skipped = _rows_to_zones(rows, coord_system, _filename_stem(file.filename))
    with closing(_db()) as conn:
        if mode == "replace":
            conn.execute("DELETE FROM zones WHERE map = 'zone'")
            base = 0
        else:
            base = conn.execute("SELECT COUNT(*) AS c FROM zones WHERE map = 'zone'").fetchone()["c"]
        created, points_total = [], 0
        for i, p in enumerate(parsed):
            cur = conn.execute(
                "INSERT INTO zones (map,name,role,color,points_json,src_points_json,coord_system,sort_order) "
                "VALUES ('zone',?,?,?,?,?,?,?)",
                (p["name"][:200], p["role"], p["color"], json.dumps(p["points"]),
                 json.dumps(p["src_points"]), coord_system, base + i),
            )
            created.append(cur.lastrowid)
            points_total += len(p["points"])
        conn.commit()
        rows_out = conn.execute(
            f"SELECT * FROM zones WHERE id IN ({','.join('?'*len(created))})", created
        ).fetchall() if created else []
    return {"zones_created": len(created), "points_total": points_total,
            "rows_skipped": skipped, "zones": [_zone_row_to_dict(r) for r in rows_out]}

@api.get("/zone/template.xlsx")
def template():
    wb = Workbook(); ws = wb.active; ws.title = "Зоны"
    ws.append(["Зона","Широта","Долгота","Цвет (необяз.)","Подпись (необяз.)"])
    for r in [["Зона 1",55.7558,37.6173,"#e53935","Штаб"],
              ["Зона 1",55.7570,37.6210,"",""],
              ["Зона 1",55.7540,37.6225,"",""],
              ["Зона 2",55.7600,37.6100,"#43a047","Объект А"],
              ["Зона 2",55.7615,37.6140,"",""]]:
        ws.append(r)
    for col, w in zip("ABCDE", (16,12,12,16,18)):
        ws.column_dimensions[col].width = w
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return Response(content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="zone_template.xlsx"'})


# ─── Прокси тайлов: Яндекс (если есть ключ) или OSM. same-origin → экспорт ────

@app.get("/tiles/{z}/{x}/{y}.png")
async def tiles(z: int, x: int, y: int):
    if not (0 <= z <= 21 and 0 <= x < (1 << z) and 0 <= y < (1 << z)):
        raise HTTPException(status_code=400, detail="Неверный тайл")
    # кеш раздельный по провайдеру, чтобы не смешивать Яндекс и OSM
    cache = TILE_CACHE / PROVIDER / str(z) / str(x) / f"{y}.png"
    if cache.exists():
        return Response(content=cache.read_bytes(), media_type="image/png",
                        headers={"Cache-Control": "public, max-age=604800"})
    if PROVIDER == "yandex":
        url = (f"{YANDEX_TILE}?l=map&x={x}&y={y}&z={z}&scale=1&lang=ru_RU&apikey={YANDEX_KEY}")
        headers = {}
        errmsg = "Не удалось получить тайл Яндекса"
    else:
        url = OSM_TILE_URL.format(z=z, x=x, y=y)
        headers = {"User-Agent": USER_AGENT}
        errmsg = "Не удалось получить тайл OSM"
    try:
        async with httpx.AsyncClient(timeout=12.0, headers=headers) as c:
            r = await c.get(url)
            r.raise_for_status()
            data = r.content
    except httpx.HTTPError:
        raise HTTPException(status_code=502, detail=errmsg)
    try:
        cache.parent.mkdir(parents=True, exist_ok=True)
        cache.write_bytes(data)
    except OSError:
        pass
    return Response(content=data, media_type="image/png",
                    headers={"Cache-Control": "public, max-age=604800"})


app.include_router(api)


@app.get("/")
def index():
    return FileResponse(STATIC_DIR / "index.html")

# Статика (JS/CSS) — после API, чтобы не перехватывать /api и /tiles.
app.mount("/", StaticFiles(directory=str(STATIC_DIR), html=True), name="static")


# ─── Запуск ───────────────────────────────────────────────────────────────────

def main():
    import uvicorn
    _init_db()
    port = int(os.environ.get("PORT", "8077"))
    host = os.environ.get("HOST", "127.0.0.1")
    url = f"http://{host if host != '0.0.0.0' else '127.0.0.1'}:{port}"
    if "--no-browser" not in sys.argv:
        threading.Timer(1.2, lambda: webbrowser.open(url)).start()
    prov = "Яндекс.Карты (ключ найден)" if PROVIDER == "yandex" else \
           "OpenStreetMap (ключ Яндекса не задан — см. README)"
    print(f"\n  Карты зон ответственности запущены: {url}"
          f"\n  Подложка: {prov}\n  (Ctrl+C — остановить)\n")
    uvicorn.run(app, host=host, port=port, log_level="warning")


if __name__ == "__main__":
    main()
