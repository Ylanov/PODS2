# app/api/v1/routers/team333_export.py
"""
Экспорт списка КОМАНДА-333 в .docx по штабному шаблону.

Документ:
  • Шапка по центру:
        РАСЧЁТ
        выделения личного состава для усиления охраны
        военного городка по сигналу «КОМАНДА-333»
        на DD.MM.YYYY
  • Таблица 5 колонок:
        Задача | Время выделения | Расчёт (по постам) | Кто выделяет | Ф.И.О.
        - «Задача» merged по всем строкам группы (вертикально)
        - «Время выделения» — берём из extra_data первой строки группы,
          merged по группе.
        - «Расчёт» — название должности (Position.name).
        - «Кто выделяет» — extra_data['deployment'].
        - «Ф.И.О.» — Slot.full_name (если пусто — пусто).
  • Подвал слева:
        Оперативный дежурный ФГКУ «ЦСООР «Лидер»
        {duty_rank}                              {duty_name}
"""
import io
from datetime import date as date_type
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from sqlalchemy.orm import Session, selectinload, joinedload

from app.models.event import Event, Group, Slot
from app.api.v1.routers.export import (
    _set_cell_border, _set_cell_shading, _cell_text, _thin_border,
)
from app.api.v1.routers.settings import get_setting


_HEADER_FILL = "D9D9D9"


def _doc_para(doc: Document, text: str, *,
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=12,
              space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def build_team333_docx(db: Session,
                       event: Event,
                       duty_rank: str,
                       duty_name: str,
                       target_date: Optional[date_type] = None) -> io.BytesIO:
    groups = (
        db.query(Group)
        .options(selectinload(Group.slots).joinedload(Slot.position))
        .filter(Group.event_id == event.id)
        .order_by(Group.order_num, Group.id)
        .all()
    )

    if target_date is None:
        target_date = event.date or date_type.today()
    date_str = target_date.strftime("%d.%m.%Y")
    org_name = get_setting(db, "org_name") or "ФГКУ «ЦСООР «Лидер»"

    # ── Документ ──────────────────────────────────────────────────────────
    doc = Document()
    section = doc.sections[0]
    section.left_margin   = Cm(1.7)
    section.right_margin  = Cm(1.7)
    section.top_margin    = Cm(1.4)
    section.bottom_margin = Cm(1.4)

    _doc_para(doc, "РАСЧЁТ",                                                    bold=True, size=14)
    _doc_para(doc, "выделения личного состава для усиления охраны",                        size=11)
    _doc_para(doc, f"военного городка по сигналу «КОМАНДА-333»",                            size=11)
    _doc_para(doc, f"на {date_str}",                                                       size=11, space_after=10)

    # Подсчёт строк таблицы: 1 шапка + sum(slots в каждой группе)
    total_slot_rows = sum(len(g.slots) for g in groups)
    table = doc.add_table(rows=1 + total_slot_rows, cols=5)
    table.autofit = False

    # Шапка
    headers = ["Задача", "Время выделения", "Расчёт (по постам)", "Кто выделяет (количество)", "Ф.И.О."]
    hdr = table.rows[0]
    for i, txt in enumerate(headers):
        _cell_text(hdr.cells[i], txt, bold=True, size_pt=10)
        _set_cell_shading(hdr.cells[i], _HEADER_FILL)
        hdr.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_border(hdr.cells[i],
                         top=_thin_border(), bottom=_thin_border(),
                         left=_thin_border(), right=_thin_border())

    # Строки + merge «Задача» / «Время» по группам
    row_idx = 1
    for group in groups:
        slots = list(group.slots)
        if not slots:
            continue
        first_extra = slots[0].get_extra() if hasattr(slots[0], 'get_extra') else {}
        task_time = first_extra.get("task_time", "")

        # Заполняем первую строку группы — она же будет merged source
        first_row = table.rows[row_idx]
        _cell_text(first_row.cells[0], group.name, bold=False, size_pt=10,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_text(first_row.cells[1], task_time, bold=False, size_pt=10)
        first_row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        first_row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # «Кто выделяет» = department + автоматический подсчёт «– N чел.»
        # по числу слотов одного управления внутри этой группы. Так
        # колонка в Word повторяет формат исходного шаблона
        # («1 упр. – 3 чел.», «Б(О) – 6 чел.», ...).
        from collections import Counter
        dept_counts = Counter(s.department or "—" for s in slots)

        # Заполняем все строки группы для колонок 2,3,4 (расчёт, кто, ФИО)
        for offset, slot in enumerate(slots):
            r = table.rows[row_idx + offset]
            position_name = slot.position.name if slot.position else ""
            dept = slot.department or "—"
            deployment    = f"{dept} – {dept_counts[dept]} чел."
            full_name     = slot.full_name or ""

            _cell_text(r.cells[2], position_name, size_pt=10,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
            _cell_text(r.cells[3], deployment, size_pt=10,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
            _cell_text(r.cells[4], full_name, size_pt=10,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
            for c in r.cells:
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _set_cell_border(c,
                                 top=_thin_border(), bottom=_thin_border(),
                                 left=_thin_border(), right=_thin_border())

        # Merge «Задача» (col 0) и «Время» (col 1) по всем строкам группы
        if len(slots) > 1:
            top_task = first_row.cells[0]
            top_time = first_row.cells[1]
            for offset in range(1, len(slots)):
                r = table.rows[row_idx + offset]
                top_task = top_task.merge(r.cells[0])
                top_time = top_time.merge(r.cells[1])

        row_idx += len(slots)

    # Ширины колонок
    widths_cm = [4.5, 2.5, 5.5, 5.0, 5.5]
    for i, w in enumerate(widths_cm):
        for r in table.rows:
            try:
                r.cells[i].width = Cm(w)
            except IndexError:
                pass

    # Подвал
    doc.add_paragraph()
    _doc_para(doc, f"Оперативный дежурный {org_name}",
              align=WD_ALIGN_PARAGRAPH.LEFT, size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p.add_run(duty_rank or "")
    run1.font.name, run1.font.size = "Times New Roman", Pt(11)
    run2 = p.add_run("\t" * 8 + (duty_name or ""))
    run2.font.name, run2.font.size = "Times New Roman", Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
