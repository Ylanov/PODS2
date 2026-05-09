# app/api/v1/routers/alert_lists.py
"""
Списки оповещения (вкладка под permission='alert_lists').

Структура:
  AlertList     — два списка (id=1, id=2). Сидируются миграцией.
  AlertPosition — словарь должностей с одним ФИО на каждую (общее для
                  всех списков). Title уникален.
  AlertSlot     — привязка должности к списку. Один слот = одна должность
                  в одном списке. UNIQUE (list_id, position_id).
  AlertMark     — отметка на день для должности. Видна в любом списке
                  где должность есть. UNIQUE (position_id, mark_date).

Что значит «общее ФИО»:
  В списке 1 и в списке 2 есть «Начальник 5 управления». Это одна и
  та же AlertPosition. Меняешь у неё primary_person_id — отображается
  в обоих списках. Ставишь V (отпуск) — то же самое.

Что свободно у каждого списка:
  Какие должности входят (можно держать в списке 1 — управления, в
  списке 2 — отделы). Порядок (sort_order у AlertSlot). Полный
  состав AlertSlot редактируется через UI («+ позиция», «📋 Шаблон»,
  drag-n-drop, «Удалить позицию»).

Эндпоинты:
  GET    /alert-lists/                                      — два списка
  GET    /alert-lists/{list_id}/slots                       — слоты списка с join'ом на position
  POST   /alert-lists/{list_id}/slots                       — добавить позицию в список
                                                              (создаёт AlertPosition если нет, AlertSlot)
  PATCH  /alert-lists/slots/{slot_id}                       — править (title/role/primary → AlertPosition;
                                                              sort_order → AlertSlot)
  DELETE /alert-lists/slots/{slot_id}                       — удалить только slot из списка
                                                              (AlertPosition останется)
  PUT    /alert-lists/{list_id}/slots/reorder               — drag-n-drop
  POST   /alert-lists/{list_id}/slots/seed                  — заполнить шаблоном
  GET    /alert-lists/{list_id}/marks?year=&month=          — отметки за месяц
  PUT    /alert-lists/slots/{slot_id}/marks/{date}          — поставить отметку
                                                              (на position, видна везде)
  DELETE /alert-lists/slots/{slot_id}/marks/{date}          — снять отметку
  GET    /alert-lists/{list_id}/export-docx?on_date=        — Word на день
  GET    /alert-lists/persons/search                        — для модалки выбора зама
"""

import logging
from datetime import date as date_type
from calendar import monthrange
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from app.api.dependencies import require_permission
from app.core.websockets import manager
from app.db.database import get_db
from app.models.alert_list import (
    AlertList, AlertPosition, AlertSlot, AlertMark,
    ALL_ALERT_MARK_TYPES, ALERT_MARK_VACATION,
    ALL_ALERT_ROLES,
)
from app.models.person import Person


logger = logging.getLogger(__name__)
router = APIRouter(dependencies=[Depends(require_permission("alert_lists"))])


# ─── Шаблон позиций по умолчанию ─────────────────────────────────────────────
DEFAULT_SLOT_TEMPLATES: list[dict] = [
    # Руководство центра
    {"title": "Первый зам. НЦ",                                "role_kind": "cnc"},
    {"title": "НШ–заместит. НЦ",                               "role_kind": "cnc"},
    {"title": "Заместитель НЦ",                                "role_kind": "cnc"},
    {"title": "Заместитель НЦ по оперативному реагированию",   "role_kind": "cnc"},
    {"title": "Заместитель НЦ по воспитательной работе",       "role_kind": "cnc"},
    {"title": "Зам. НЦ по тылу",                               "role_kind": "cnc"},
    {"title": "Зам. НЦ по вооружению",                         "role_kind": "cnc"},
    {"title": "Зам. НШ",                                       "role_kind": "cnc"},
    {"title": "Зам. НШ (по орг.-моб. раб.)",                   "role_kind": "cnc"},
    {"title": "зам. НШ по оперативной работе",                 "role_kind": "cnc"},
    # Управления
    {"title": "1 Управление",                                  "role_kind": "upr"},
    {"title": "2 Управление",                                  "role_kind": "upr"},
    {"title": "3 Управление",                                  "role_kind": "upr"},
    {"title": "4 Управление",                                  "role_kind": "upr"},
    {"title": "5 Управление",                                  "role_kind": "upr"},
    {"title": "6 Управление",                                  "role_kind": "upr"},
    {"title": "7 Управление",                                  "role_kind": "upr"},
    {"title": "8 Управление",                                  "role_kind": "upr"},
    # Отделы и группы
    {"title": "Отдел кадров",                                  "role_kind": "otd"},
    {"title": "Отдел воспитательной работы",                   "role_kind": "otd"},
    {"title": "Отдел организационный и комплектования",        "role_kind": "otd"},
    {"title": "Отдел эксплуатации зданий",                     "role_kind": "otd"},
    {"title": "Отдел (профессиональной подготовки)",           "role_kind": "otd"},
    {"title": "Отдел (организации контрактной работы)",        "role_kind": "otd"},
    {"title": "Нач. отд. – гл. бухгалтер",                     "role_kind": "otd"},
    {"title": "Начальник отдела-нач. связи",                   "role_kind": "otd"},
    {"title": "Начальник клуба",                               "role_kind": "otd"},
    {"title": "Начальник группы-комендант",                    "role_kind": "otd"},
    # Службы
    {"title": "Юридическая служба",                            "role_kind": "otd"},
    {"title": "Психологическая служба",                        "role_kind": "otd"},
    {"title": "Вещевая служба",                                "role_kind": "otd"},
    {"title": "Продовольственная служба",                      "role_kind": "otd"},
    {"title": "Автомобильная служба",                          "role_kind": "otd"},
    {"title": "Инженерная служба",                             "role_kind": "otd"},
    {"title": "Воздушно-десантная служба",                     "role_kind": "otd"},
    {"title": "Служба горючего и смазочных материалов",        "role_kind": "otd"},
    {"title": "Служба артиллерийского вооружения",             "role_kind": "otd"},
    {"title": "Служба защиты государственной тайны",           "role_kind": "otd"},
    {"title": "Служба ППЗ и СР",                               "role_kind": "otd"},
    {"title": "Служба РХБЗ",                                   "role_kind": "otd"},
    # Прочее
    {"title": "ВАИ",                                           "role_kind": "otd"},
    {"title": "БАЗА (ОБЕСПЕЧЕНИЯ)",                            "role_kind": "otd"},
    {"title": "Оркестр - Военный дирижер",                     "role_kind": "otd"},
]


# ─── Pydantic ────────────────────────────────────────────────────────────────

class _PersonRef(BaseModel):
    id:        int
    full_name: str
    rank:      Optional[str] = None
    position_title: Optional[str] = None


class ListOut(BaseModel):
    id:   int
    name: str


class SlotOut(BaseModel):
    """
    Slot отдаётся фронту с разворотом полей AlertPosition в плоскую
    структуру — UI как был, ничего менять не нужно. position_id отдаём
    отдельно, потому что фронт может пригодиться (для marks-логики).
    """
    id:          int
    list_id:     int
    position_id: int
    title:       str
    role_kind:   str
    sort_order:  int
    primary_person: Optional[_PersonRef] = None


class SlotIn(BaseModel):
    title:             str = Field(..., min_length=1, max_length=200)
    role_kind:         str = Field(default="upr")
    sort_order:        int = 0
    primary_person_id: Optional[int] = None


class SlotPatch(BaseModel):
    title:             Optional[str] = Field(default=None, min_length=1, max_length=200)
    role_kind:         Optional[str] = None
    sort_order:        Optional[int] = None
    primary_person_id: Optional[int] = None
    primary_person_id_set: bool = False


class MarkOut(BaseModel):
    """
    slot_id оставлен в выдаче для бэк-совместимости фронта (он строит
    Map по этому ключу). Бэкенд по slot_id → position_id и работает.

    source:
      'manual' — отметка введена вручную в списках оповещения (AlertMark)
      'duty'   — отметка взята из графика нарядов (DutyMark) по primary_person
                  данной должности. На фронте такие — read-only, нельзя
                  снять кликом «снять» (они подтянутся обратно).
    """
    slot_id:   int
    mark_date: date_type
    mark_type: str
    substitute_person: Optional[_PersonRef] = None
    source:    str = "manual"
    duty_schedule_title: Optional[str] = None   # для подсказки «откуда»


class MarkIn(BaseModel):
    mark_type:           str
    substitute_person_id: Optional[int] = None


class ReorderPayload(BaseModel):
    slot_ids: List[int]


class SeedPayload(BaseModel):
    pass


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _person_ref(p: Optional[Person]) -> Optional[_PersonRef]:
    if not p:
        return None
    return _PersonRef(
        id=p.id, full_name=p.full_name, rank=p.rank, position_title=p.position_title,
    )


def _slot_out(s: AlertSlot) -> SlotOut:
    pos = s.position
    return SlotOut(
        id=s.id,
        list_id=s.list_id,
        position_id=pos.id,
        title=pos.title,
        role_kind=pos.role_kind,
        sort_order=s.sort_order,
        primary_person=_person_ref(pos.primary_person),
    )


def _mark_out(slot_id: int, m: AlertMark) -> MarkOut:
    return MarkOut(
        slot_id=slot_id,
        mark_date=m.mark_date,
        mark_type=m.mark_type,
        substitute_person=_person_ref(m.substitute_person),
        source="manual",
    )


# Приоритет «полосовых» отметок: чем недоступнее — тем выше. Используется
# при выборе одной отметки на день когда у человека несколько DutyMark за
# один день в разных графиках.
_DUTY_MARK_PRIORITY = {"V": 4, "T": 3, "H": 2, "N": 1}


def _get_or_create_position(db: Session, title: str, role_kind: str,
                            primary_person_id: Optional[int]) -> AlertPosition:
    """
    Найти AlertPosition по title или создать новую. role_kind /
    primary_person_id применяются только при создании; при существующей —
    не перезаписываем (если нужно изменить — через PATCH /slots/{id}).
    """
    title = title.strip()
    pos = db.query(AlertPosition).filter(AlertPosition.title == title).first()
    if pos:
        return pos
    pos = AlertPosition(
        title=title,
        role_kind=role_kind,
        primary_person_id=primary_person_id,
    )
    db.add(pos)
    db.flush()
    return pos


async def _sync_person_position_title(db: Session, position: AlertPosition) -> Optional[int]:
    """
    Обратная связь со справочником людей: когда к должности привязали
    конкретного Person — переписываем у него Person.position_title в
    соответствие с AlertPosition.title. Так в Базе людей у Иванова сразу
    видно «Начальник 5 управления», а не пустое поле.

    Возвращает person_id если был апдейт (для WS broadcast'а), иначе None.

    Не очищаем title если primary_person_id стал None — пользователь мог
    сам прописать должность вручную, не будем затирать молча.
    """
    if not position.primary_person_id:
        return None
    person = db.query(Person).filter(Person.id == position.primary_person_id).first()
    if not person:
        return None
    new_title = position.title
    if person.position_title == new_title:
        return None
    person.position_title = new_title
    db.flush()
    return person.id


# ─── Lists ───────────────────────────────────────────────────────────────────

@router.get("/", response_model=List[ListOut], summary="Два списка оповещения")
def list_lists(db: Session = Depends(get_db)):
    rows = db.query(AlertList).order_by(AlertList.id.asc()).all()
    return [ListOut(id=r.id, name=r.name) for r in rows]


# ─── Slots CRUD ──────────────────────────────────────────────────────────────

@router.get("/{list_id}/slots", response_model=List[SlotOut],
            summary="Слоты выбранного списка")
def list_slots(list_id: int, db: Session = Depends(get_db)):
    if not db.query(AlertList).filter(AlertList.id == list_id).first():
        raise HTTPException(status_code=404, detail="Список не найден")
    rows = (
        db.query(AlertSlot)
        .filter(AlertSlot.list_id == list_id)
        .order_by(AlertSlot.sort_order.asc(), AlertSlot.id.asc())
        .all()
    )
    return [_slot_out(s) for s in rows]


@router.post("/{list_id}/slots", response_model=SlotOut, status_code=201,
             summary="Добавить позицию в список")
async def create_slot(list_id: int, payload: SlotIn, db: Session = Depends(get_db)):
    if not db.query(AlertList).filter(AlertList.id == list_id).first():
        raise HTTPException(status_code=404, detail="Список не найден")
    if payload.role_kind not in ALL_ALERT_ROLES:
        raise HTTPException(status_code=400, detail=f"role_kind должен быть один из {ALL_ALERT_ROLES}")
    if payload.primary_person_id is not None:
        if not db.query(Person).filter(Person.id == payload.primary_person_id).first():
            raise HTTPException(status_code=400, detail="primary_person_id не найден")

    pos = _get_or_create_position(db, payload.title, payload.role_kind, payload.primary_person_id)
    # Если такая позиция уже есть в этом списке — отказ.
    exists = (
        db.query(AlertSlot)
        .filter(AlertSlot.list_id == list_id, AlertSlot.position_id == pos.id)
        .first()
    )
    if exists:
        raise HTTPException(status_code=409, detail="Эта должность уже есть в списке")

    slot = AlertSlot(list_id=list_id, position_id=pos.id, sort_order=payload.sort_order or 0)
    db.add(slot)
    # Если при создании сразу указали primary_person_id — синхронизируем
    # должность в Базе людей.
    synced_person_id = await _sync_person_position_title(db, pos)
    db.commit()
    db.refresh(slot)
    await manager.broadcast({"action": "alert_lists_update"})
    if synced_person_id:
        await manager.broadcast({"action": "person_update", "person_id": synced_person_id})
    return _slot_out(slot)


@router.patch("/slots/{slot_id}", response_model=SlotOut,
              summary="Изменить позицию (title/role/primary — общее; sort_order — у слота)")
async def patch_slot(slot_id: int, payload: SlotPatch, db: Session = Depends(get_db)):
    s = db.query(AlertSlot).filter(AlertSlot.id == slot_id).first()
    if not s:
        raise HTTPException(status_code=404, detail="Слот не найден")
    pos = s.position

    # title/role_kind/primary_person_id — это поля AlertPosition. Изменения
    # отразятся во всех списках где эта позиция присутствует.
    if payload.title is not None:
        new_title = payload.title.strip()
        if new_title != pos.title:
            # Если такой title уже есть у другой позиции — нельзя
            # (нарушит UNIQUE). UI должен валидировать заранее.
            other = db.query(AlertPosition).filter(
                AlertPosition.title == new_title,
                AlertPosition.id    != pos.id,
            ).first()
            if other:
                raise HTTPException(
                    status_code=409,
                    detail="Должность с таким названием уже существует. "
                           "Используйте её или придумайте другое имя.",
                )
            pos.title = new_title
    if payload.role_kind is not None:
        if payload.role_kind not in ALL_ALERT_ROLES:
            raise HTTPException(status_code=400, detail=f"role_kind должен быть один из {ALL_ALERT_ROLES}")
        pos.role_kind = payload.role_kind
    if payload.primary_person_id_set:
        if payload.primary_person_id is not None:
            if not db.query(Person).filter(Person.id == payload.primary_person_id).first():
                raise HTTPException(status_code=400, detail="primary_person_id не найден")
        pos.primary_person_id = payload.primary_person_id

    # sort_order — у самого AlertSlot, локально для списка.
    if payload.sort_order is not None:
        s.sort_order = payload.sort_order

    # Синхронизируем должность в Базе людей: если поменяли primary_person_id
    # ИЛИ title — у привязанного Person обновим position_title.
    synced_person_id = None
    if payload.primary_person_id_set or payload.title is not None:
        synced_person_id = await _sync_person_position_title(db, pos)

    db.commit()
    db.refresh(s)
    await manager.broadcast({"action": "alert_lists_update"})
    if synced_person_id:
        await manager.broadcast({"action": "person_update", "person_id": synced_person_id})
    return _slot_out(s)


@router.delete("/slots/{slot_id}", status_code=204,
               summary="Удалить позицию из списка (AlertPosition не трогается)")
async def delete_slot(slot_id: int, db: Session = Depends(get_db)):
    s = db.query(AlertSlot).filter(AlertSlot.id == slot_id).first()
    if not s:
        raise HTTPException(status_code=404, detail="Слот не найден")
    db.delete(s)
    db.commit()
    await manager.broadcast({"action": "alert_lists_update"})


@router.put("/{list_id}/slots/reorder", summary="Переупорядочить слоты списка (drag-n-drop)")
async def reorder_slots(list_id: int, payload: ReorderPayload, db: Session = Depends(get_db)):
    if not db.query(AlertList).filter(AlertList.id == list_id).first():
        raise HTTPException(status_code=404, detail="Список не найден")
    rows = (
        db.query(AlertSlot)
        .filter(AlertSlot.list_id == list_id, AlertSlot.id.in_(payload.slot_ids))
        .all()
    )
    by_id = {s.id: s for s in rows}
    for idx, sid in enumerate(payload.slot_ids):
        s = by_id.get(sid)
        if s:
            s.sort_order = idx
    db.commit()
    await manager.broadcast({"action": "alert_lists_update", "list_id": list_id})
    return {"updated": len(by_id)}


# ─── Шаблон ──────────────────────────────────────────────────────────────────

@router.get("/template/preview", summary="Шаблон стандартных позиций")
def template_preview():
    return DEFAULT_SLOT_TEMPLATES


@router.post("/sync-persons-positions",
             summary="Синхронизировать должности всех AlertPosition в Базу людей")
async def sync_persons_positions(db: Session = Depends(get_db)):
    """
    Проходит по всем AlertPosition с primary_person_id и переписывает у
    привязанных Person'ов position_title в соответствие с title должности.
    Полезно для разовой починки данных, когда часть привязок делалась
    до new-77 (когда обратной записи в Person ещё не было).

    Возвращает {updated, skipped} — сколько Person'ов реально изменилось
    и сколько уже было синхронизировано.
    """
    rows = (
        db.query(AlertPosition)
        .filter(AlertPosition.primary_person_id.isnot(None))
        .all()
    )
    updated_ids: list[int] = []
    skipped = 0
    for pos in rows:
        person = db.query(Person).filter(Person.id == pos.primary_person_id).first()
        if not person:
            continue
        if person.position_title == pos.title:
            skipped += 1
            continue
        person.position_title = pos.title
        updated_ids.append(person.id)
    db.commit()

    # Один общий broadcast для всех изменённых; персональные не критичны.
    if updated_ids:
        await manager.broadcast({"action": "alert_lists_update"})
        for pid in updated_ids:
            await manager.broadcast({"action": "person_update", "person_id": pid})

    return {"updated": len(updated_ids), "skipped": skipped, "total_with_person": len(rows)}


@router.post("/{list_id}/slots/seed",
             summary="Заполнить список стандартными позициями (idempotent)")
async def seed_slots(list_id: int, payload: SeedPayload, db: Session = Depends(get_db)):
    """
    Для каждой записи шаблона:
      • если AlertPosition с таким title нет — создаём
      • если в этом списке этой position_id ещё нет — добавляем AlertSlot
      • иначе — пропускаем

    sort_order у новых слотов продолжается от текущего max.
    """
    if not db.query(AlertList).filter(AlertList.id == list_id).first():
        raise HTTPException(status_code=404, detail="Список не найден")

    existing_position_ids = {
        row[0] for row in
        db.query(AlertSlot.position_id).filter(AlertSlot.list_id == list_id).all()
    }
    max_order = (
        db.query(AlertSlot.sort_order)
        .filter(AlertSlot.list_id == list_id)
        .order_by(AlertSlot.sort_order.desc())
        .first()
    )
    next_order = (max_order[0] + 1) if max_order else 0

    created = 0
    skipped = 0
    for tpl in DEFAULT_SLOT_TEMPLATES:
        pos = _get_or_create_position(db, tpl["title"], tpl["role_kind"], None)
        if pos.id in existing_position_ids:
            skipped += 1
            continue
        slot = AlertSlot(list_id=list_id, position_id=pos.id, sort_order=next_order)
        db.add(slot)
        existing_position_ids.add(pos.id)
        next_order += 1
        created += 1

    db.commit()
    await manager.broadcast({"action": "alert_lists_update", "list_id": list_id})
    return {"created": created, "skipped": skipped, "total": created + skipped}


# ─── Marks ───────────────────────────────────────────────────────────────────

@router.get("/{list_id}/marks", response_model=List[MarkOut],
            summary="Все отметки списка за месяц (manual + derived из DutyMark)")
def list_marks(
    list_id: int,
    year:    int = Query(..., ge=2000, le=2100),
    month:   int = Query(..., ge=1, le=12),
    db:      Session = Depends(get_db),
):
    """
    Возвращает union:
      • ручные AlertMark для позиций этого списка (source='manual')
      • derived из DutyMark за тот же месяц для primary_person каждой
        позиции (source='duty'); типы N/V/T/H — наряд/отпуск/командировка/
        госпиталь. Тип U (увольнение/выходной) НЕ переносим — он не
        блокирует оповещение.

    Приоритет на одну ячейку: manual > derived. Если у одного человека
    несколько DutyMark за день в разных графиках — выбирается самая
    «недоступная» (V > T > H > N).
    """
    from app.models.duty import DutyMark, DutySchedule

    if not db.query(AlertList).filter(AlertList.id == list_id).first():
        raise HTTPException(status_code=404, detail="Список не найден")
    last  = monthrange(year, month)[1]
    start = date_type(year, month, 1)
    end   = date_type(year, month, last)

    # 1. Ручные отметки.
    manual_rows = (
        db.query(AlertSlot.id, AlertMark)
        .join(AlertMark, AlertMark.position_id == AlertSlot.position_id)
        .filter(
            AlertSlot.list_id == list_id,
            AlertMark.mark_date >= start,
            AlertMark.mark_date <= end,
        )
        .all()
    )

    out: list[MarkOut] = []
    seen_keys: set[tuple[int, date_type]] = set()   # (slot_id, date) уже занятые ручной

    for slot_id, m in manual_rows:
        out.append(_mark_out(slot_id, m))
        seen_keys.add((slot_id, m.mark_date))

    # 2. Derived из DutyMark. Для каждого слота списка с привязанной person
    #    тащим её отметки за период.
    slot_rows = (
        db.query(AlertSlot.id, AlertPosition.primary_person_id)
        .join(AlertPosition, AlertSlot.position_id == AlertPosition.id)
        .filter(
            AlertSlot.list_id == list_id,
            AlertPosition.primary_person_id.isnot(None),
        )
        .all()
    )
    person_to_slots: dict[int, list[int]] = {}
    for slot_id, person_id in slot_rows:
        person_to_slots.setdefault(person_id, []).append(slot_id)

    if person_to_slots:
        duty_rows = (
            db.query(DutyMark, DutySchedule.title)
            .join(DutySchedule, DutyMark.schedule_id == DutySchedule.id)
            .filter(
                DutyMark.person_id.in_(person_to_slots.keys()),
                DutyMark.mark_type.in_(["N", "V", "T", "H"]),
                DutyMark.duty_date >= start,
                DutyMark.duty_date <= end,
            )
            .all()
        )

        # Для каждой пары (person_id, date) — выбираем ОДНУ отметку с
        # наивысшим приоритетом, если в нескольких графиках разные.
        best_per_person_day: dict[tuple[int, date_type], tuple[DutyMark, str]] = {}
        for dm, sched_title in duty_rows:
            key = (dm.person_id, dm.duty_date)
            cur = best_per_person_day.get(key)
            if not cur or _DUTY_MARK_PRIORITY.get(dm.mark_type, 0) > _DUTY_MARK_PRIORITY.get(cur[0].mark_type, 0):
                best_per_person_day[key] = (dm, sched_title)

        # Раскладываем по слотам, пропуская уже занятые ручной отметкой.
        for (person_id, day), (dm, sched_title) in best_per_person_day.items():
            for slot_id in person_to_slots.get(person_id, []):
                if (slot_id, day) in seen_keys:
                    continue
                out.append(MarkOut(
                    slot_id=slot_id,
                    mark_date=day,
                    mark_type=dm.mark_type,
                    substitute_person=None,
                    source="duty",
                    duty_schedule_title=sched_title,
                ))
                seen_keys.add((slot_id, day))

    return out


@router.put("/slots/{slot_id}/marks/{mark_date}", response_model=MarkOut,
            summary="Поставить/обновить отметку (на ДОЛЖНОСТЬ — видна в обоих списках)")
async def upsert_mark(
    slot_id:   int,
    mark_date: date_type,
    payload:   MarkIn,
    db:        Session = Depends(get_db),
):
    slot = db.query(AlertSlot).filter(AlertSlot.id == slot_id).first()
    if not slot:
        raise HTTPException(status_code=404, detail="Слот не найден")
    position_id = slot.position_id

    if payload.mark_type not in ALL_ALERT_MARK_TYPES:
        raise HTTPException(status_code=400, detail=f"mark_type должен быть один из {ALL_ALERT_MARK_TYPES}")

    if payload.mark_type == ALERT_MARK_VACATION:
        if not payload.substitute_person_id:
            raise HTTPException(
                status_code=400,
                detail="Для отпуска (V) обязательно указать заместителя — substitute_person_id",
            )
        if not db.query(Person).filter(Person.id == payload.substitute_person_id).first():
            raise HTTPException(status_code=400, detail="substitute_person_id не найден")
        sub_id = payload.substitute_person_id
    else:
        sub_id = None

    existing = (
        db.query(AlertMark)
        .filter(AlertMark.position_id == position_id, AlertMark.mark_date == mark_date)
        .first()
    )
    if existing:
        existing.mark_type = payload.mark_type
        existing.substitute_person_id = sub_id
        mark = existing
    else:
        mark = AlertMark(
            position_id=position_id,
            mark_date=mark_date,
            mark_type=payload.mark_type,
            substitute_person_id=sub_id,
        )
        db.add(mark)
    db.commit()
    db.refresh(mark)
    await manager.broadcast({"action": "alert_lists_update"})
    return _mark_out(slot_id, mark)


@router.delete("/slots/{slot_id}/marks/{mark_date}", status_code=204,
               summary="Снять отметку")
async def delete_mark(
    slot_id:   int,
    mark_date: date_type,
    db:        Session = Depends(get_db),
):
    slot = db.query(AlertSlot).filter(AlertSlot.id == slot_id).first()
    if not slot:
        return
    mark = (
        db.query(AlertMark)
        .filter(AlertMark.position_id == slot.position_id,
                AlertMark.mark_date  == mark_date)
        .first()
    )
    if not mark:
        return
    db.delete(mark)
    db.commit()
    await manager.broadcast({"action": "alert_lists_update"})


# ─── Экспорт в Word на конкретный день ───────────────────────────────────────

@router.get("/{list_id}/export-docx", summary="Экспорт списка на конкретный день в .docx")
def export_alert_list_docx(
    list_id: int,
    on_date: date_type = Query(..., description="День, на который формируется список"),
    db:      Session = Depends(get_db),
):
    from io import BytesIO
    from urllib.parse import quote
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    alert_list = db.query(AlertList).filter(AlertList.id == list_id).first()
    if not alert_list:
        raise HTTPException(status_code=404, detail="Список не найден")

    slots = (
        db.query(AlertSlot)
        .filter(AlertSlot.list_id == list_id)
        .order_by(AlertSlot.sort_order.asc(), AlertSlot.id.asc())
        .all()
    )

    position_ids = [s.position_id for s in slots]
    # Карта position_id → (mark_type, substitute_person | None) с приоритетом
    # manual > derived. Сначала собираем manual (могут перекрыть derived).
    marks_by_position: dict[int, tuple[str, Optional[Person]]] = {}

    if position_ids:
        manual_marks = (
            db.query(AlertMark)
            .filter(
                AlertMark.position_id.in_(position_ids),
                AlertMark.mark_date == on_date,
            )
            .all()
        )
        for m in manual_marks:
            marks_by_position[m.position_id] = (m.mark_type, m.substitute_person)

        # Derived для тех позиций, по которым ещё нет manual.
        from app.models.duty import DutyMark
        person_ids = [s.position.primary_person_id for s in slots
                      if s.position.primary_person_id]
        if person_ids:
            duty_today = (
                db.query(DutyMark)
                .filter(
                    DutyMark.person_id.in_(person_ids),
                    DutyMark.mark_type.in_(["N", "V", "T", "H"]),
                    DutyMark.duty_date == on_date,
                )
                .all()
            )
            best: dict[int, DutyMark] = {}
            for dm in duty_today:
                cur = best.get(dm.person_id)
                if not cur or _DUTY_MARK_PRIORITY.get(dm.mark_type, 0) > _DUTY_MARK_PRIORITY.get(cur.mark_type, 0):
                    best[dm.person_id] = dm
            for s in slots:
                pid = s.position.primary_person_id
                if pid and pid in best and s.position_id not in marks_by_position:
                    marks_by_position[s.position_id] = (best[pid].mark_type, None)

    MARK_TITLES = {
        "N": "Наряд", "O": "Ответственный",
        "V": "Отпуск", "T": "Командировка", "H": "Госпиталь",
    }

    doc = Document()
    section = doc.sections[0]
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{alert_list.name}\nна {on_date.strftime('%d.%m.%Y')}")
    run.bold = True
    run.font.size = Pt(14)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Должность"
    hdr[2].text = "ФИО"
    hdr[3].text = "Отметка"
    for c in hdr:
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True

    for idx, slot in enumerate(slots, start=1):
        pos = slot.position
        entry = marks_by_position.get(pos.id)
        mark_type = entry[0] if entry else None
        sub_person = entry[1] if entry else None

        who = None
        suffix = ""
        if mark_type == "V" and sub_person:
            who = sub_person
            suffix = " (замещает)"
        elif pos.primary_person:
            who = pos.primary_person
        full_name = who.full_name if who else "—"
        rank      = (who.rank + " ") if who and who.rank else ""

        mark_label = MARK_TITLES.get(mark_type, "") if mark_type else ""

        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = pos.title
        row[2].text = f"{rank}{full_name}{suffix}"
        row[3].text = mark_label

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{alert_list.name.replace(' ', '_')}_{on_date.strftime('%Y-%m-%d')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition":
                f"attachment; filename=\"{filename}\"; filename*=UTF-8''{quote(filename)}",
        },
    )


# ─── Поиск кандидатов на зама ────────────────────────────────────────────────

@router.get("/persons/search", response_model=List[_PersonRef],
            summary="Поиск Person для модалки выбора (primary / зам)")
def search_persons(
    q:    str = Query("", max_length=200),
    role: Optional[str] = Query(default=None, description="upr/otd/cnc — для подсказки, не фильтр"),
    root: Optional[str] = Query(default=None, max_length=200,
                                description="корень должности — РЕЗУЛЬТАТЫ С ИМ ВЫШЕ, но не отсекаем"),
    db:   Session = Depends(get_db),
):
    """
    Возвращает активных людей, отсортированных так:
      1. Совпадение position_title по `root` (если задан) — наверху;
      2. Дальше — по ФИО.

    Раньше для upr/otd фильтр был ЖЁСТКИМ (position_title ILIKE %root%),
    из-за чего у юзера с пустыми position_title в Базе людей выпадашка
    оказывалась пустой и выбрать вообще никого было нельзя. Теперь это
    только сортировка, не отсечение — пользователь видит всех и может
    выбрать. Когда position_title заполнятся (через alert_lists их
    автозаполняет new-77), сортировка сама собой станет полезной.
    """
    qry = db.query(Person).filter(Person.fired_at.is_(None))
    if q:
        qry = qry.filter(Person.full_name.ilike(f"%{q}%"))
    rows = qry.order_by(Person.full_name.asc()).limit(80).all()

    # Сортировка-приоритет: совпавшие по root наверх. role=cnc — без
    # boost'а, для центра ищем кого угодно.
    if role in ("upr", "otd") and root:
        root_l = root.lower()
        rows.sort(key=lambda p: (
            0 if (p.position_title or "").lower().find(root_l) != -1 else 1,
            p.full_name.lower(),
        ))

    return [_person_ref(p) for p in rows if p]
