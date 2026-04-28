# app/api/v1/routers/groza_export.py
"""
Экспорт списка ГРОЗА-555 в .docx по штабному шаблону.

Документ соответствует образцу:
  • Шапка по центру:
        Состав сил и средств ФГКУ «ЦСООР «Лидер»
        по сигналу «ГРОЗА-555»
        на DD.MM.YYYY
  • Основная таблица: 6 колонок (№ п/п | Воинское звание | ФИО |
    № документа | Должность, техника | Подразделение). Группы — строки
    с merged-cell с названием группы (фон серый).
  • При наличии групп с is_supplementary=True — заголовок-разделитель
        Состав сил и средств обеспечения доставки
        в район сбора (до аэропорта погрузки)
    и вторая таблица той же структуры.
  • Подвал слева:
        Оперативный дежурный ФГКУ «ЦСООР «Лидер»
        {duty_rank}                   {duty_name}

    duty_rank/duty_name — параметры запроса (фронт показывает диалог
    с дефолтами из /settings: duty_rank, duty_name).

Автодетект ГРОЗА-режима: события с хотя бы одной is_supplementary-группой
выгружаются через этот builder вместо стандартного export.py.
"""

import io
from datetime import date as date_type, datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from sqlalchemy.orm import Session, selectinload

from app.models.event import Event, Group, Slot, Position
from app.api.v1.routers.export import (
    _set_cell_border, _set_cell_shading, _cell_text, _thin_border,
)
from app.api.v1.routers.settings import get_setting


_HEADER_FILL = "D9D9D9"   # серый — шапка таблицы и group-rows


def _doc_para(doc: Document, text: str, *,
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=12,
              space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def _build_table(doc: Document, slots_by_group: list[tuple[str, list]]):
    """
    slots_by_group: [(group_name, [slot, ...]), ...]
    Все слоты в этой таблице. Группа = одна объединённая строка с названием.
    """
    if not slots_by_group:
        return

    total_rows = sum(1 + len(slots) for _, slots in slots_by_group) + 1   # +1 шапка
    table = doc.add_table(rows=total_rows, cols=6)
    table.autofit = False

    # ── Шапка таблицы ──────────────────────────────────────────────────────
    hdr = table.rows[0]
    headers = ["№ п/п", "Воинское звание", "Фамилия Имя Отчество",
               "№ документа", "Должность, техника", "Подразделение"]
    for i, txt in enumerate(headers):
        _cell_text(hdr.cells[i], txt, bold=True, size_pt=10)
        _set_cell_shading(hdr.cells[i], _HEADER_FILL)
        hdr.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_border(hdr.cells[i],
                         top=_thin_border(), bottom=_thin_border(),
                         left=_thin_border(), right=_thin_border())

    # ── Группы и слоты ─────────────────────────────────────────────────────
    row_idx = 1
    person_idx = 0
    for group_name, slots in slots_by_group:
        # Заголовок группы — объединённая строка
        gh = table.rows[row_idx]
        merged = gh.cells[0]
        for c in gh.cells[1:]:
            merged = merged.merge(c)
        _cell_text(merged, group_name, bold=True, size_pt=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(merged, _HEADER_FILL)
        for c in gh.cells:
            _set_cell_border(c,
                             top=_thin_border(), bottom=_thin_border(),
                             left=_thin_border(), right=_thin_border())
        row_idx += 1

        # Слоты группы
        for slot in slots:
            person_idx += 1
            row = table.rows[row_idx]
            extra = slot.get_extra()
            position_name = slot.position.name if slot.position else ""

            cells_data = [
                str(person_idx),
                slot.rank or "",
                slot.full_name or "",
                slot.doc_number or "",
                position_name,
                extra.get("subdivision", ""),
            ]
            for i, txt in enumerate(cells_data):
                _cell_text(row.cells[i], txt, size_pt=10,
                           align=WD_ALIGN_PARAGRAPH.LEFT if i in (1, 2, 4, 5) else WD_ALIGN_PARAGRAPH.CENTER)
                row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _set_cell_border(row.cells[i],
                                 top=_thin_border(), bottom=_thin_border(),
                                 left=_thin_border(), right=_thin_border())
            row_idx += 1

    # ── Ширины колонок ─────────────────────────────────────────────────────
    widths_cm = [1.1, 2.6, 5.0, 2.6, 5.5, 2.4]
    for i, w in enumerate(widths_cm):
        for r in table.rows:
            try:
                r.cells[i].width = Cm(w)
            except IndexError:
                pass


def _is_groza_event(event: Event) -> bool:
    """Событие выгружается в ГРОЗА-формате если есть supplementary-группы."""
    return any(getattr(g, "is_supplementary", False) for g in event.groups)


def build_groza_docx(db: Session,
                     event: Event,
                     duty_rank: str,
                     duty_name: str,
                     target_date: Optional[date_type] = None) -> io.BytesIO:
    """
    Собирает .docx документ ГРОЗА-555 для конкретного развернутого списка.
    target_date — дата сверху (на DD.MM.YYYY). Если не передана, берём event.date,
    иначе сегодня.
    """
    # Загружаем группы со слотами
    groups = (
        db.query(Group)
        .options(selectinload(Group.slots).joinedload(Slot.position))
        .filter(Group.event_id == event.id)
        .order_by(Group.order_num, Group.id)
        .all()
    )
    main_groups = [(g.name, list(g.slots)) for g in groups if not getattr(g, "is_supplementary", False)]
    aux_groups  = [(g.name, list(g.slots)) for g in groups if     getattr(g, "is_supplementary", False)]

    # Дата
    if target_date is None:
        target_date = event.date or date_type.today()
    date_str = target_date.strftime("%d.%m.%Y")

    org_name = get_setting(db, "org_name") or "ФГКУ «ЦСООР «Лидер»"

    # ── Документ ──────────────────────────────────────────────────────────
    doc = Document()
    section = doc.sections[0]
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    # Шапка
    _doc_para(doc, f"Состав сил и средств {org_name}", size=12)
    _doc_para(doc, "по сигналу «ГРОЗА-555»",            size=12, bold=True)
    _doc_para(doc, f"на {date_str}",                    size=12, space_after=10)

    # Основная таблица
    _build_table(doc, main_groups)

    # Дополнительный список — заголовок и вторая таблица
    if aux_groups:
        doc.add_paragraph()   # отступ
        _doc_para(doc, "Состав сил и средств обеспечения доставки", size=12)
        _doc_para(doc, "в район сбора (до аэропорта погрузки)",     size=12, space_after=8)
        _build_table(doc, aux_groups)

    # Подвал — оперативный дежурный
    doc.add_paragraph()
    _doc_para(doc, f"Оперативный дежурный {org_name}",
              align=WD_ALIGN_PARAGRAPH.LEFT, size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p.add_run(duty_rank or "")
    run1.font.name, run1.font.size = "Times New Roman", Pt(11)
    run2 = p.add_run("\t" * 8 + (duty_name or ""))
    run2.font.name, run2.font.size = "Times New Roman", Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
