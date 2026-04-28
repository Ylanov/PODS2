# app/api/v1/routers/duty_export.py
"""
Экспорт графика наряда в .docx по шаблону штаба.

Документ:
  • альбомная ориентация A4
  • справа сверху — блок «УТВЕРЖДАЮ» (тексты из /settings:
    print_approve_position/rank/name)
  • заголовок «Г Р А Ф И К» по центру
  • подзаголовок: «{schedule.title} на {month} {year} года»
  • таблица: № п/п | В/звание Ф.И.О. | 1 | 2 | … | 31
      - чёрные ячейки   — наряды (Н)
      - бледно-зелёные  — отпуска (V)
      - бледно-янтарные — увольнения (У)
      - синие           — резервы (РЗ)
      - серые           — выходные (Сб/Вс/праздники)
  • подпись внизу (print_footer_*)

Используется и admin-роутером и dept-роутером — общий код в
build_duty_schedule_docx(...).
"""

import io
from calendar import monthrange
from datetime import date as date_type
from typing import List

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from sqlalchemy.orm import Session, joinedload

from app.models.duty import (
    DutySchedule, DutySchedulePerson, DutyMark,
    MARK_DUTY, MARK_LEAVE, MARK_VACATION, MARK_RESERVE,
)
from app.api.v1.routers.export import (
    _set_cell_border, _set_cell_shading, _cell_text, _thin_border,
)
from app.api.v1.routers.settings import get_setting


_MONTHS_RU = [
    "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль",   "Август",  "Сентябрь", "Октябрь", "Ноябрь", "Декабрь",
]

# Цветовая палитра ячеек. Hex без '#'. Подобрана под Word — печатается
# адекватно как в цвете, так и в ч/б (градации серого различимы).
_FILL_DUTY     = "1E293B"   # тёмный графит — наряд
_FILL_RESERVE  = "3B82F6"   # синий — резерв
_FILL_LEAVE    = "D97706"   # янтарь — увольнение
_FILL_VACATION = "9DD9C5"   # бледно-зелёный — отпуск
_FILL_WEEKEND  = "F1F5F9"   # светло-серый — выходные
_FILL_HEADER   = "E5E7EB"   # серый — шапка таблицы

_RANK_ORDER = [
    "Генерал армии", "Генерал-полковник", "Генерал-лейтенант", "Генерал-майор",
    "Полковник", "Подполковник", "Майор",
    "Капитан", "Старший лейтенант", "Лейтенант", "Младший лейтенант",
    "Старший прапорщик", "Прапорщик",
    "Старшина", "Старший сержант", "Сержант", "Младший сержант",
    "Ефрейтор", "Рядовой",
]
_RANK_INDEX = {r.lower(): i for i, r in enumerate(_RANK_ORDER)}


def _rank_key(rank: str | None) -> int:
    """Индекс звания — старшие сверху. Нестандартные/пустые в конец."""
    if not rank:
        return len(_RANK_ORDER) + 1
    return _RANK_INDEX.get(rank.strip().lower(), len(_RANK_ORDER))


def _approve_block(doc: Document, settings: dict, year: int, month: int):
    """Блок «УТВЕРЖДАЮ» — выровненный вправо, без таблицы."""
    pos  = settings.get("print_approve_position", "")
    rank = settings.get("print_approve_rank",     "")
    name = settings.get("print_approve_name",     "")

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_after = Pt(2)

    def _line(text: str, bold: bool = False):
        run = para.add_run(text + "\n")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = bold

    _line("УТВЕРЖДАЮ", bold=True)
    if pos:  _line(pos)
    if rank: _line(rank)
    _line(f'«____» {month:02d}.{year} г.    {name}'.rstrip())


def _title_block(doc: Document, schedule_title: str, year: int, month: int):
    """«Г Р А Ф И К / {название} на {месяц} {год} года»."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after  = Pt(2)
    run = title.add_run("Г Р А Ф И К")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(8)
    run = sub.add_run(f"{schedule_title or ''} на {_MONTHS_RU[month - 1]} {year} года")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def _footer_block(doc: Document, settings: dict):
    pos  = settings.get("print_footer_position", "")
    rank = settings.get("print_footer_rank",     "")
    name = settings.get("print_footer_name",     "")
    if not (pos or rank or name):
        return
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    if pos:
        run = para.add_run(pos + "\n")
        run.font.name, run.font.size = "Times New Roman", Pt(11)
    if rank:
        run = para.add_run(rank + "\n")
        run.font.name, run.font.size = "Times New Roman", Pt(11)
    if name:
        run = para.add_run(f"_____________ {name}")
        run.font.name, run.font.size = "Times New Roman", Pt(11)


def _build_table(doc: Document,
                 persons: List[dict],
                 marks_by_person: dict,
                 year: int, month: int):
    days_count = monthrange(year, month)[1]
    # Колонки: №п/п, ФИО+звание, 1..days_count
    cols = 2 + days_count
    table = doc.add_table(rows=1 + len(persons), cols=cols)
    table.autofit = False

    # ── Шапка ──────────────────────────────────────────────────────────────
    hdr = table.rows[0]
    _cell_text(hdr.cells[0], "№ п/п",         bold=True, size_pt=9)
    _cell_text(hdr.cells[1], "В / звание Ф.И.О.", bold=True, size_pt=9,
               align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_shading(hdr.cells[0], _FILL_HEADER)
    _set_cell_shading(hdr.cells[1], _FILL_HEADER)

    for i in range(days_count):
        day = i + 1
        cell = hdr.cells[2 + i]
        dow  = date_type(year, month, day).weekday()  # 0=Пн … 6=Вс
        is_weekend = dow >= 5
        _cell_text(cell, str(day), bold=True, size_pt=8)
        _set_cell_shading(cell, _FILL_WEEKEND if is_weekend else _FILL_HEADER)

    for c in hdr.cells:
        _set_cell_border(c, top=_thin_border(), bottom=_thin_border(),
                            left=_thin_border(), right=_thin_border())

    # ── Строки людей ───────────────────────────────────────────────────────
    for r_idx, p in enumerate(persons, start=1):
        row = table.rows[r_idx]
        _cell_text(row.cells[0], str(r_idx), size_pt=9)
        # ФИО + звание (звание в начале)
        rank_prefix = (p["rank"] + " ") if p["rank"] else ""
        _cell_text(row.cells[1], f'{rank_prefix}{p["full_name"]}',
                   size_pt=9, align=WD_ALIGN_PARAGRAPH.LEFT)

        person_marks = marks_by_person.get(p["person_id"], {})
        for i in range(days_count):
            day = i + 1
            cell = row.cells[2 + i]
            iso  = date_type(year, month, day).isoformat()
            mt   = person_marks.get(iso)

            dow = date_type(year, month, day).weekday()
            is_weekend = dow >= 5

            text = ""
            color = None
            if mt == MARK_DUTY:
                color, text = _FILL_DUTY,     "Н"
            elif mt == MARK_RESERVE:
                color, text = _FILL_RESERVE,  "РЗ"
            elif mt == MARK_LEAVE:
                color, text = _FILL_LEAVE,    "У"
            elif mt == MARK_VACATION:
                color, text = _FILL_VACATION, "О"
            elif is_weekend:
                color = _FILL_WEEKEND

            _cell_text(cell, text, size_pt=8, bold=bool(text))
            if color:
                _set_cell_shading(cell, color)

            # Текст белым на тёмных фонах для читаемости.
            if mt in (MARK_DUTY, MARK_RESERVE, MARK_LEAVE):
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for c in row.cells:
            _set_cell_border(c, top=_thin_border(), bottom=_thin_border(),
                                left=_thin_border(), right=_thin_border())

    # ── Ширины колонок ─────────────────────────────────────────────────────
    table.columns[0].width = Cm(0.9)    # №
    table.columns[1].width = Cm(6.0)    # ФИО
    day_col_w = Cm(0.7)
    for i in range(days_count):
        table.columns[2 + i].width = day_col_w


def build_duty_schedule_docx(db: Session,
                             schedule: DutySchedule,
                             year: int,
                             month: int) -> io.BytesIO:
    """
    Собирает .docx документ для конкретного графика и месяца.
    Возвращает BytesIO-буфер с готовым файлом.
    """
    # Загружаем людей и марки одним заходом
    persons_q = (
        db.query(DutySchedulePerson)
        .options(joinedload(DutySchedulePerson.person))
        .filter(DutySchedulePerson.schedule_id == schedule.id)
        .all()
    )
    persons = [
        {
            "person_id": p.person.id,
            "full_name": p.person.full_name,
            "rank":      p.person.rank,
        }
        for p in persons_q
    ]
    persons.sort(key=lambda p: (_rank_key(p["rank"]), (p["full_name"] or "").lower()))

    last_day = monthrange(year, month)[1]
    marks = (
        db.query(DutyMark)
        .filter(
            DutyMark.schedule_id == schedule.id,
            DutyMark.duty_date   >= date_type(year, month, 1),
            DutyMark.duty_date   <= date_type(year, month, last_day),
        )
        .all()
    )
    marks_by_person: dict[int, dict[str, str]] = {}
    for m in marks:
        marks_by_person.setdefault(m.person_id, {})[m.duty_date.isoformat()] = m.mark_type

    # Settings для шапки/подписи
    settings_keys = (
        "print_approve_position", "print_approve_rank", "print_approve_name",
        "print_footer_position",  "print_footer_rank",  "print_footer_name",
    )
    settings = {k: get_setting(db, k) for k in settings_keys}

    # ── Документ ───────────────────────────────────────────────────────────
    doc = Document()
    section = doc.sections[0]
    section.orientation   = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.4)
    section.top_margin  = section.bottom_margin = Cm(1.2)

    _approve_block(doc, settings, year, month)
    _title_block(doc, schedule.title or "", year, month)
    _build_table(doc, persons, marks_by_person, year, month)
    _footer_block(doc, settings)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
